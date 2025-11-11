# 版本信息
__version__ = "1.2"

import os
import re
import sys
import json
import shutil
import openpyxl
import zipfile
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from PyQt5.QtWidgets import QMessageBox, QApplication, QFileDialog, QWidget, QVBoxLayout, QHBoxLayout, QTextEdit, QLineEdit, QPushButton, QSizePolicy, QMainWindow, QLineEdit, QLabel, QDialog, QColorDialog, QCheckBox
from PyQt5.QtCore import QObject, pyqtSignal, Qt, QFile
from PyQt5.QtGui import QIcon

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):  # PyInstaller打包后的临时目录
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

def load_stylesheet(stylesheet_path):
    absolute_path = resource_path(stylesheet_path)
    with open(absolute_path, "r", encoding="utf-8") as file:
        return file.read()    
    
# 报错及提示输出包装函数

# 报错提示
def print_ten_pentagram():
    print('★' * 10)
    return

def print_info_path_finding(path):
    print(f" --> 正在搜索目录【{path}】 ... 请勿中止！")
    return

def print_info_file_finding(file_path):
    print(f" --> 正在搜索文件【{file_path}】 ... 请勿中止！")
    return

def print_info_path_found(path):
    print(f" --> 已搜索到目录【{path}】！")
    return

def print_info_file_found(file_path):
    print(f" --> 已搜索到文件【{file_path}】！")
    return

def print_error_path_found(path):
    print(f" 【Error】 未搜索到目录【{path}】！")
    return

def print_error_file_found(file_path):
    print(f" 【Error】 未搜索到文件【{file_path}】！")
    return

def print_error_not_path(path):
    print(f" 【Error】 【{path}】不是合法目录！")
    return

def print_error_not_file(file_path):
    print(f" 【Error】 【{file_path}】不是合法文件！")
    return

def print_info_path_building(path):
    print(f" --> 正在创建目录【{path}】 ... 请勿中止！")
    return

def print_info_file_building(file_path):
    print(f" --> 正在创建文件【{file_path}】 ... 请勿中止！")
    return

def print_info_path_build(path):
    print(f" --> 目录【{path}】创建成功！")
    return

def print_info_file_build(file_path):
    print(f" --> 文件【{file_path}】创建成功！")
    return

def print_error_path_build(path):
    print(f" 【Error】 目录【{path}】创建失败！")
    return

def print_error_file_build(file_path):
    print(f" 【Error】 文件【{file_path}】创建失败！")
    return

def print_info_config_space_build(config_space):
    print(f" --> 已搜索到配置目录！文件夹位置：【{config_space}】")
    return

def print_error_config_space_build(config_space, e):
    print(f" 【Error】 未搜索到配置目录！文件夹位置：【{config_space}】错误详情：{e}")
    return

def print_info_config_space_startup(config_path, information_excel_document_path):
    print(f" --> 配置目录初始化成功！配置文件位置：【{config_path}】信息表位置：【{information_excel_document_path}】")
    return

def print_error_config_space_startup(config_space, e):
    print(f" 【Error】 配置目录初始化失败！文件夹位置：【{config_space}】错误详情：{e}")
    return

def print_info_find_target_numeric_folder(work_space, target_numeric_folder_path):
    print(f" --> 在【{work_space}】中搜索到当前月文件夹！文件夹位置：【{target_numeric_folder_path}】")
    return

def print_error_find_target_numeric_folder(e):
    print(f" 【Error】 未搜索到当前月文件夹！错误详情：{e}")
    return

def print_info_find_detail_word_document(target_numeric_folder_path, detail_word_document_path):
    print(f" --> 在【{target_numeric_folder_path}】中搜索到报销明细 Word 文档！文件位置：【{detail_word_document_path}】")
    return

def print_error_find_detail_word_document(e):
    print(f" 【Error】 未搜索到报销明细 Word 文档！错误详情：{e}")
    return

def print_info_find_detail_excel_document(target_numeric_folder_path, detail_excel_document_path):
    print(f" --> 在【{target_numeric_folder_path}】中搜索到报销明细 Excel 文档！文件位置：【{detail_excel_document_path}】")
    return

def print_error_find_detail_excel_document(e):
    print(f" 【Error】 未搜索到报销明细 Excel 文档！错误详情：{e}")
    return

def print_info_find_explanation_document(document_type, target_numeric_folder_path, explanation_document_path):
    print(f" --> 在【{target_numeric_folder_path}】中搜索到{document_type}情况说明文档！文件位置：【{explanation_document_path}】")
    return

def print_error_find_explanation_document(document_type, e):
    print(f" 【Error】 未搜索到{document_type}情况说明文档！错误详情：{e}")
    return

def print_info_initialize_detail_word_document_title1(detail_word_document_path):
    print(f" --> 报销明细 Word 文档正在初始化！进度：1/3 ... 请勿中止！文件位置：【{detail_word_document_path}】")
    return

def print_info_initialize_detail_word_document_title2(detail_word_document_path):
    print(f" --> 报销明细 Word 文档正在初始化！进度：2/3 ... 请勿中止！文件位置：【{detail_word_document_path}】")
    return

def print_info_initialize_detail_word_document_table(detail_word_document_path):
    print(f" --> 报销明细 Word 文档正在初始化！进度：3/3 ... 请勿中止！文件位置：【{detail_word_document_path}】")
    return

def print_info_initialize_detail_word_document(detail_word_document_path):
    print(f" --> 报销明细 Word 文档初始化完成！文件位置：【{detail_word_document_path}】")
    return

def print_info_initialize_detail_word_document_finished(detail_word_document_path):
    print(f" --> 已存在初始化完成的报销明细 Word 文档！文件位置：【{detail_word_document_path}】")
    return

def print_error_initialize_detail_word_document(detail_word_document_path, e):
    print(f" 【Error】 报销明细 Word 文档初始化失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_initialize_detail_excel_document(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档初始化完成！文件位置：【{detail_excel_document_path}】")
    return

def print_error_initialize_detail_excel_document(detail_excel_document_path, e):
    print(f" 【Error】 报销明细 Excel 文档初始化失败！文件位置：【{detail_excel_document_path}】错误详情：{e}")
    return

def print_info_delete_sum_detail_word_document(detail_word_document_path):
    print(f" --> 报销明细 Word 文档数据清洗完成！文件位置：【{detail_word_document_path}】")
    return

def print_error_delete_sum_detail_word_document(detail_word_document_path, e):
    print(f" 【Error】 报销明细 Word 文档数据清洗失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_detail_word_document_add_table(student_name, use_purpose, money_amount):
    print(f" --> 【{student_name}】【{use_purpose}】【{money_amount}】写入成功！")
    return

def print_error_detail_word_document_add_table(detail_word_document_path, e):
    print(f" 【Error】 数据写入失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_sort_detail_word_document(detail_word_document_path):
    print(f" --> 报销明细 Word 文档整理完成！文件位置：【{detail_word_document_path}】")
    return

def print_error_sort_detail_word_document(detail_word_document_path, e):
    print(f" 【Error】 报销明细 Word 文档整理失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_update_detail_word_document_suffix(detail_word_document_path, updated_detail_word_document_path):
    print(f" --> 报销明细 Word 文档后缀日期更新完成！【{detail_word_document_path}】已重命名为【{updated_detail_word_document_path}】！")
    return

def print_error_update_detail_word_document_suffix(detail_word_document_path, e):
    print(f" 【Error】 报销明细 Word 文档后缀日期更新失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_read_detail_word_document(detail_word_document_path):
    print(f" --> 报销明细 Word 文档读取完成！文件位置：【{detail_word_document_path}】")
    return

def print_error_read_detail_word_document(detail_word_document_path, e):
    print(f" 【Error】 报销明细 Word 文档读取失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_write_detail_excel_document(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档正在写入！进度：1/6 ... 请勿中止！文件位置：【{detail_excel_document_path}】")
    return

def print_error_write_detail_excel_document(detail_excel_document_path, e):
    print(f" 【Error】 报销明细 Excel 文档写入失败！文件位置：【{detail_excel_document_path}】错误详情：{e}")
    return

def print_info_calculate_detail_excel_document_personal(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档正在计算！进度：2/6 ... 请勿中止！文件位置：【{detail_excel_document_path}】")
    return

def print_info_calculate_detail_excel_document_sum(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档正在计算！进度：3/6 ... 请勿中止！文件位置：【{detail_excel_document_path}】")
    return

def print_info_calculate_detail_excel_document_second_table(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档正在计算！进度：4/6 ... 请勿中止！文件位置：【{detail_excel_document_path}】")
    return

def print_info_calculate_detail_excel_document_second_table_sum(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档正在计算！进度：5/6 ... 请勿中止！文件位置：【{detail_excel_document_path}】")
    return

def print_error_calculate_detail_excel_document(detail_excel_document_path, e):
    print(f" 【Error】 报销明细 Excel 文档计算失败！文件位置：【{detail_excel_document_path}】错误详情：{e}")
    return

def print_info_beautify_detail_excel_document(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档正在渲染！进度：6/6 ... 请勿中止！文件位置：【{detail_excel_document_path}】")
    return

def print_error_beautify_detail_excel_document(detail_excel_document_path, e):
    print(f" 【Error】 报销明细 Excel 文档渲染失败！文件位置：【{detail_excel_document_path}】错误详情：{e}")
    return

def print_info_detail_excel_document_finished(detail_excel_document_path):
    print(f" --> 报销明细 Excel 文档已生成！文件位置：【{detail_excel_document_path}】")
    return

def print_info_back_detail_word_document(detail_word_document_path, initial_datas_dict):
    print(f" --> 报销明细 Word 文档回写完成！文件位置：【{detail_word_document_path}】")
    print(f" --> 总计：{initial_datas_dict['initial_sum']} 元")
    return

def print_error_back_detail_word_document(detail_word_document_path, e):
    print(f" 【Error】 报销明细 Word 文档回写失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_explanation_document_add_title(document_type, explanation_document_path):
    print(f" --> {document_type}情况说明文档正在写入！进度：1/3 ... 请勿中止！文件位置：【{explanation_document_path}】")
    return

def print_info_explanation_document_add_paragraph(document_type, explanation_document_path):
    print(f" --> {document_type}情况说明文档正在写入！进度：2/3 ... 请勿中止！文件位置：【{explanation_document_path}】")
    return

def print_info_explanation_document_add_inscribed(document_type, explanation_document_path):
    print(f" --> {document_type}情况说明文档正在写入！进度：3/3 ... 请勿中止！文件位置：【{explanation_document_path}】")
    return

def print_info_write_explanation_document(document_type, explanation_document_path):
    print(f" --> {document_type}情况说明文档写入完成！文件位置：【{explanation_document_path}】")
    return

def print_error_write_explanation_document(document_type, explanation_document_path, e):
    print(f" 【Error】 {document_type}情况说明文档写入失败！文件位置：【{explanation_document_path}】错误详情：{e}")
    return

def print_error_no_destination():
    print(" 【Error】 请输入出租车目的地！")
    return

def print_info_student_name_with_id(student_name, student_id):
    print(f" --> {student_name}的学号是：{student_id}！")
    return

def print_error_student_name_with_id(student_name, e):
    print(f" 【Error】 {student_name}的学号未找到！错误详情：{e}")
    return

def print_info_output_documents(output_space, zip_folder_path):
    print(f" --> 输出报销文档到【{output_space}】成功！文件位置：【{zip_folder_path}】")
    return

def print_error_output_documents(output_space, e):
    print(f" 【Error】 输出报销文档到【{output_space}】失败！错误详情：{e}")
    return

def print_error_special_material_input():
    print(" 【Error】 请成对输入特殊材料及其用途！")
    return

def print_info_open_detail_word_document(detail_word_document_path):
    print(f" --> 报销明细 Word 文档打开成功！文件位置：【{detail_word_document_path}】")
    return

def print_error_open_detail_word_document(detail_word_document_path, e):
    print(f" 【Error】 报销明细 Word 文档打开失败！文件位置：【{detail_word_document_path}】错误详情：{e}")
    return

def print_info_save_student_id_in_information_excel_document(information_excel_document_path, student_name, student_id):
    print(f" --> 【{student_name}】【{student_id}】已成功保存到信息表！文件位置：【{information_excel_document_path}】")
    return

def print_error_save_student_id_in_information_excel_document(information_excel_document_path, e):
    print(f" 【Error】 信息保存失败！文件位置：【{information_excel_document_path}】错误详情：{e}")
    return

def print_info_check_new_work_space(new_work_space):
    print(f" --> 新工作目录【{new_work_space}】符合要求！")
    return

def print_error_check_new_work_space(new_work_space, e):
    print(f" 【Error】 新工作目录【{new_work_space}】不符合要求！错误详情：{e}")
    return

def print_info_to_new_work_space(folder_path, new_folder_path):
    print(f" --> 成功将【{folder_path}】移动至【{new_folder_path}】！")
    return

def print_info_move_to_new_work_space(new_work_space):
    print(f" --> 新工作目录【{new_work_space}】准备完成！")
    return

def print_error_move_to_new_work_space(new_work_space, e):
    print(f" 【Error】 新工作目录【{new_work_space}】准备失败！错误详情：{e}")
    return

def print_info_check(student_name, use_purpose, money_amount):
    print(f" --> 【{student_name}】【{use_purpose}】【{money_amount}】已录入！")
    return

def print_info_find_numeric_folders(work_space, numeric_folders):
    for numeric_folder in numeric_folders:
        print(f" --> 在工作目录【{work_space}】下找到【{numeric_folder}】！")
    return

def print_error_find_numeric_folders(work_space, e):
    print(f" 【Error】 在工作目录【{work_space}】下未找到符合要求的文件夹！错误详情：{e}")
    return

def print_info_search(student_name, all_datas, collected_datas):
    print(f" --> 查询成功！")
    print('--------------------')
    print(f"【{student_name}】的报销记录：")
    print('--------------------')
    for file_data in all_datas:
        if file_data.get('data'):
            print(f"【{file_data.get('year_month')}】")
            for row_data in file_data.get('data'):
                print(f"【{row_data[0]}】【{row_data[1]}】【{row_data[2]}】")
            print('--------------------')
    print(f"总计：")
    for i in range(6):
        print(f"【{use_purpose_number_dict[i]}】【{collected_datas[use_purpose_number_dict[i]]}】")
    print('--------------------')
    print(f" --> 检索完毕！")
    return

def print_error_search(e):
    print(f" 【Error】 查询失败！错误详情：{e}")
    return

def print_info_del_numeric_folder(numeric_folder_path):
    print(f" --> 删除文件夹【{numeric_folder_path}】成功！")
    return

def print_error_del_numeric_folder(numeric_folder_path, e):
    print(f" 【Error】 删除文件夹【{numeric_folder_path}】失败！错误详情：{e}")
    return

def print_info_restart_work_space(work_space):
    print(f" --> 重置工作目录【{work_space}】成功！")
    return

def print_info_restart_config_space(config_space):
    print(f" --> 重置配置文件夹【{config_space}】成功！")
    return

def print_error_restart_config_space(config_space, e):
    print(f" 【Error】 重置配置文件夹【{config_space}】失败！错误详情：{e}")
    return

def main_text_append_error_input(self, input):
    self.main_text.append(f" 【Error】 请正确输入{input}！\n")
    return

def main_text_append_error_notfound_search_student_name(self):
    self.main_text.append(" 【Error】 查无此人！\n")
    return

def main_text_append_error_notfound_files(self):
    self.main_text.append(" 【Error】 查无数据！\n")
    return

def main_text_append_info_setting(self, setting_type, setting):
    self.parent().main_text.append(f" --> 【{setting_type}】：{setting} 已保存！\n")
    return

def warning_input(self, input):
    QMessageBox.warning(self, "警告", f"请正确输入{input}！")
    return

#第〇个功能小函数开始

# 配置文件默认值
default_config = {
    "work_space": r"C:\Easy Accounting",
    "output_space": os.path.join(os.path.expanduser("~"), "Desktop"),
    "usr_name": "",
    "padding": False,
    "colors": ['#C7EDCC', '#FFE2D2', '#FAF9DE', '#CC99CC', '#87CEFA', '#ECECEF']
}

# 用途字典
use_purpose_number_dict = {
    0: '材料',
    1: '市内交通',
    2: '物流',
    3: '打印',
    4: '差旅',
    5: '论文'
}

# 配置文件属性修改
def update_config(config_path, key, value):
    # 尝试加载现有配置文件
    if os.path.exists(config_path):
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
    else:
        config = {}

    # 更新指定的配置项
    config[key] = value

    # 将更新后的配置写回到文件中
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=4)

# 初始化配置文件（存在则不初始化，不存在则初始化）
def startup_config_path(config_path):
    
    # 检查路径中的文件是否存在
    if os.path.exists(config_path):
        # 尝试加载现有配置文件
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
    else:
        # 如果文件不存在，则初始化配置并写入
        config = default_config.copy()

        # 创建文件并写入默认配置
        os.makedirs(os.path.dirname(config_path), exist_ok=True)  # 如果路径不存在，创建目录
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=4)

    # 返回配置（无论是加载的还是新创建的）
    return config

# 初始化信息表（存在则不初始化，不存在则初始化）
def startup_information_excel_document_path(information_excel_document_path):
    # 如果文件不存在，则创建新文件并初始化数据
    if not os.path.exists(information_excel_document_path):
        # 创建一个新的工作簿
        wb = openpyxl.Workbook()
        ws = wb.active  # 获取当前活动的工作表

        # 设置标题行
        ws.cell(row=2, column=2, value="姓名")  # 第二行第二列为“姓名”
        ws.cell(row=2, column=3, value="学号")  # 第二行第三列为“学号”

        # 从外部文件读取学生数据
        student_data_file = resource_path("students_data.txt")
        student_data = []
        
        try:
            with open(student_data_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line:  # 跳过空行
                        parts = line.split('\t')
                        if len(parts) == 2:  # 确保每行有姓名和学号
                            student_data.append(tuple(parts))
        except FileNotFoundError:
            print(f"错误：未找到学生数据文件 students_data.txt")
            print(f"请参考 students_data.txt.example 创建该文件")
            print(f"文件格式：每行一个学生，姓名和学号用制表符分隔")
            raise
        except Exception as e:
            print(f"错误：读取学生数据文件时出错：{str(e)}")
            raise

        # 填入数据（从第三行开始，第二列为姓名，第三列为学号）
        for index, (student_name, student_id) in enumerate(student_data, start=3):
            ws.cell(row=index, column=2, value=student_name.strip())
            ws.cell(row=index, column=3, value=student_id.strip())

        # 设置对齐样式：水平和垂直居中
        alignment_style = Alignment(horizontal='center', vertical='center')

        # 设置字体样式
        font_style = Font(name='等线', size=11, bold=False)
   
        # 遍历所有行和列
        for row in ws.iter_rows(min_row=1, min_col=1):
            # 设置行高为16
            ws.row_dimensions[row[0].row].height = 16
            for cell in row:
                if cell.value is not None:  # 如果单元格有内容
                    # 应用对齐样式
                    cell.alignment = alignment_style
                    # 应用字体样式
                    cell.font = font_style

        # 设置列宽
        column_widths = {
            'A': 3,
            'B': 12,
            'C': 20,
        }
        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

        # 保存工作簿
        wb.save(information_excel_document_path)
        wb.close()

#第〇个功能小函数结束

#第〇个功能大函数开始
def startup(config_space, config_path, information_excel_document_path):
    try:
        ensure_path_exist(config_space)
        print_info_config_space_build(config_space)
    except Exception as e:
        print_ten_pentagram()
        print_error_config_space_build(config_space, e)
        return
    
    try:
        startup_config_path(config_path)
        startup_information_excel_document_path(information_excel_document_path)
        print_info_config_space_startup(config_path, information_excel_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_config_space_startup(config_space, e)
        return
    
#第〇个功能大函数结束

# 第一个功能小函数开始
# ensure_path_exist(path)
# find_target_numeric_folder(work_space)
# find_detail_word_document(target_folder_path)

# find_student_id_by_name(information_excel_document_path, student_name)

# detail_word_document_add_title(doc, text)
# detail_word_document_table_borders_set(table)
# detail_word_document_table_cell_style_set(cell)
# detail_word_document_add_table(doc, student_name, use_purpose, money_amount)
# initialize_detail_word_document(detail_word_document_path, information_excel_document_path, usr_name)

# 确保指定路径存在
# 如果路径无效（例如不是路径格式）则抛出 ValueError
# 如果路径有效但不存在，则创建该路径
# 如果路径已存在，则不进行任何操作
def ensure_path_exist(path):
    
    print_info_path_finding(path)

    # 检查路径是否是一个有效路径格式
    if not isinstance(path, str) or not path.strip():
        print_ten_pentagram()
        print_error_not_path(path)
        raise ValueError
    
    # 判断路径是否存在
    if os.path.exists(path):
        if not os.path.isdir(path):
            print_ten_pentagram()
            print_error_not_path(path)
            raise ValueError
        # 路径存在且是一个目录，无需操作
        return path
    
    # 路径不存在，创建路径
    try:
        print_info_path_building(path)
        os.makedirs(path)
        print_info_path_build(path)
    except Exception as e:
        print_ten_pentagram()
        print_error_path_build(path)
        raise OSError
    
    return path

# 找数字文件夹202409
def find_target_numeric_folder(work_space, create=False):

    # 搜索目录
    try:
        ensure_path_exist(work_space)
        print_info_path_found(work_space)
    except Exception as e:
        print_ten_pentagram()
        print_error_path_found(work_space)
        raise FileNotFoundError

    # 列出所有文件夹，筛选出数字命名的
    subfolders = [folder for folder in os.listdir(work_space) if os.path.isdir(os.path.join(work_space, folder))]
    numeric_folders = [folder for folder in subfolders if folder.isdigit()]
    
    if numeric_folders and not create:
        # 选择数字最大的文件夹
        target_numeric_folder = max(numeric_folders)
    else:
        # 创建当前年月为名称的文件夹
        current_year_month = datetime.now().strftime("%Y%m")
        target_numeric_folder = current_year_month
        os.makedirs(os.path.join(work_space, target_numeric_folder))

    # 获取目标文件夹的绝对路径
    target_numeric_folder_path = os.path.join(work_space, target_numeric_folder)
    return target_numeric_folder_path

# 找当月报销明细 Word 文档   yes：选择第一个（不能删除）；no：创建一个
def find_detail_word_document(target_folder_path):
    # 搜索目录
    try:
        ensure_path_exist(target_folder_path)
        print_info_path_found(target_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_path_found(target_folder_path)
        raise FileNotFoundError

    detail_word_documents = [file for file in os.listdir(target_folder_path) if file.endswith(".docx") and "报销明细" in file]
    
    if detail_word_documents:
        # 若存在符合条件的文档，选择第一个
        detail_word_document = detail_word_documents[0]
    else:
        # 若不存在，则创建新的 Word 文档
        current_year_month = datetime.now().strftime("%Y年") + str(int(datetime.now().strftime("%m"))) + "月"
        current_year_month_date_num = datetime.now().strftime("%Y%m%d")
        detail_word_document = f"{current_year_month}报销明细_{current_year_month_date_num}.docx"
        # 创建并保存空白文档
        doc = Document()
        doc.save(os.path.join(target_folder_path, detail_word_document))

    # 返回报销明细 Word 文档的绝对路径
    detail_word_document_path = os.path.join(target_folder_path, detail_word_document)
    return detail_word_document_path

# 依据姓名在信息表中找学号
def find_student_id_by_name(information_excel_document_path, student_name):
    try:
        # 加载 Excel 文件
        workbook = openpyxl.load_workbook(information_excel_document_path)
        sheet = workbook.active  # 默认使用第一个工作表

        # 遍历 Excel 的每一行
        for row in sheet.iter_rows(min_row=3):  # 从第3行开始
            name_cell = row[1]  # 第2列：学生姓名
            id_cell = row[2]    # 第3列：学生学号

            # 检查姓名是否匹配
            if name_cell.value == student_name:
                student_id = str(id_cell.value)
                workbook.save(information_excel_document_path)
                workbook.close()
                return student_id  # 返回对应的学号
        workbook.save(information_excel_document_path)
        workbook.close()
        return None
    except Exception as e:
        workbook.save(information_excel_document_path)
        workbook.close()
        return None

# 报销明细 Word 文档加标题行
def detail_word_document_add_title(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph()
    # 水平居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 设置段前/段后间距
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1  # 单倍行距
    paragraph_format.space_before = Pt(12)  # 段前0.5行 (0.5 * 24磅)
    paragraph_format.space_after = Pt(12)  # 段后0.5行

    # 设置字号
    run = paragraph.add_run(text)
    run.font.size = Pt(16)  # 小二字体（16磅）

    # 设置中文为宋体，英文及数字为新罗马
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.append(rFonts)

# 报销明细 Word 文档表格线型设置
def detail_word_document_table_borders_set(table):
    tbl = table._tbl

    # 确保命名空间与实际文档中的匹配
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    tblPr = tbl.find(".//w:tblPr", namespaces=namespaces)

    if tblPr is None:
        # 如果 tblPr 不存在，则创建
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    # 创建表格边框元素
    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # 实线
        border.set(qn('w:sz'), '4')  # 边框宽度，0.5磅 = 4 (twip单位)
        border.set(qn('w:space'), '0')  # 无间距
        border.set(qn('w:color'), '000000')  # 黑色
        borders.append(border)

    tblPr.append(borders)

# 报销明细 Word 文档表格单元格格式设置
def detail_word_document_table_cell_style_set(cell):

    # 设置单元格文字格式
    for paragraph in cell.paragraphs:
        # 水平居中
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run = paragraph.runs[0]
        # 设置字号
        run.font.size = Pt(16)
        # 设置中文为宋体，英文及数字为新罗马
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rPr.append(rFonts)

        # 设置段落间距
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1  # 单倍行距
        paragraph_format.space_before = Pt(6)  # 段前0.5行（1行=12pt）
        paragraph_format.space_after = Pt(6)  # 段后0.5行

    # 设置单元格内容垂直居中
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)  # 确保附加的元素为 XML 元素类型

# 报销明细 Word 文档加表格行
def detail_word_document_add_table(doc, student_name, use_purpose, money_amount):

    # 确保所有输入转换为字符串
    student_name = str(student_name)
    use_purpose = str(use_purpose)
    money_amount = str(money_amount)

    # 获取文档中的表格列表
    tables = doc.tables
    
    if tables:
        # 如果已有表格，向第一个表格添加新行
        table = tables[0]
        row = table.add_row()
    else:
        # 如果没有表格，创建一个新的表格
        table = doc.add_table(rows=1, cols=3)
        # 设置表格边框样式
        detail_word_document_table_borders_set(table)
        # 设置表格居中
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row = table.rows[0]

    # 设置列宽
    for i, cell in enumerate(row.cells):
        if i == 0:  # 第一列（姓名列）
            cell.width = Inches(1.2)  # 设置第一列宽度
        elif i == 1:  # 第二列（用途列）
            cell.width = Inches(3)  # 设置第二列宽度
        else:  # 第三列（金额列）
            cell.width = Inches(1.5)  # 设置第三列宽度

    # 添加内容到表格的每个单元格
    for idx, text in enumerate([student_name, use_purpose, money_amount]):
        cell = row.cells[idx]
        cell.text = text
        detail_word_document_table_cell_style_set(cell)

    return

# 初始化报销明细 Word 文档
def initialize_detail_word_document(detail_word_document_path, information_excel_document_path, usr_name):
    
    doc = Document(detail_word_document_path)

    # 获取文档中的表格列表（检查是否已被初始化）
    tables = doc.tables
    if tables:
        # 保存文档
        doc.save(detail_word_document_path)
        print_info_initialize_detail_word_document_finished(detail_word_document_path)
        return

    usr_id = find_student_id_by_name(information_excel_document_path, usr_name)

    detail_word_document_add_title(doc, datetime.now().strftime("%Y年") + str(int(datetime.now().strftime("%m"))) + "月报销明细")
    print_info_initialize_detail_word_document_title1(detail_word_document_path)

    detail_word_document_add_title(doc, "负责人：" + usr_name + "      " + "学号：" + usr_id)
    print_info_initialize_detail_word_document_title2(detail_word_document_path)

    detail_word_document_add_table(doc, "姓名", "用途", "金额")
    print_info_initialize_detail_word_document_table(detail_word_document_path)

    # 保存文档
    doc.save(detail_word_document_path)

    return
# 第一个功能小函数结束
# ensure_path_exist(path)
# find_target_numeric_folder(work_space)
# find_detail_word_document(target_folder_path)

# find_student_id_by_name(information_excel_document_path, student_name)

# detail_word_document_add_title(doc, text)
# detail_word_document_table_borders_set(table)
# detail_word_document_table_cell_style_set(cell)
# detail_word_document_add_table(doc, student_name, use_purpose, money_amount)
# initialize_detail_word_document(detail_word_document_path, information_excel_document_path, usr_name)


# 第一个功能大函数开始
# 准备好报销明细 Word 文档
def create_detail_word_document(work_space, information_excel_document_path, usr_name):

    # 搜索当前月文件夹
    try:
        target_numeric_folder_path = find_target_numeric_folder(work_space, True)
        print_info_find_target_numeric_folder(work_space, target_numeric_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_target_numeric_folder(e)
        return

    # 搜索报销明细 Word 文档
    try:
        detail_word_document_path = find_detail_word_document(target_numeric_folder_path)
        print_info_find_detail_word_document(target_numeric_folder_path, detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_detail_word_document(e)
        return

    # 初始化报销明细 Word 文档
    try:
        initialize_detail_word_document(detail_word_document_path, information_excel_document_path, usr_name)
        print_info_initialize_detail_word_document(detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_initialize_detail_word_document(detail_word_document_path, e)
        return

    return detail_word_document_path
# 第一个功能大函数结束

# 第二个功能小函数开始
# use_purpose_order_switch(use_purpose)
# sort_detail_word_document(detail_word_document_path, usr_name)
# update_detail_word_document_suffix(detail_word_document_path)

# 用途排序 switch
def use_purpose_order_switch(use_purpose):
    use_purpose_order_dict = {
        '出租车费': 0,
        '过路费': 1,
        '过桥费': 2,
        '地铁费': 3,
        '寄件运费': 4,
        '飞机票': 5,
        '高铁票': 6,
        '住宿费': 7,
        '餐饮费': 8,
        '论文费': 9,
        '打印费': 10,
    }
    return use_purpose_order_dict.get(use_purpose, 11)

# 重排报销明细 Word 文档数据
def sort_detail_word_document(detail_word_document_path, usr_name):

    # 打开Word文档
    doc = Document(detail_word_document_path)
    
    # 获取文档中的表格列表
    tables = doc.tables
    if tables:
        table = tables[0]
    
        # 初始化数组
        detail_word_document_datas = []
        # 初始化集合（去重）
        student_name_datas = []
        # 遍历表格的每一行，从第二行开始（跳过表头）
        for row_idx, row in enumerate(table.rows[1:], start=2):
            # 获取当前行的三个单元格数据
            student_name = str(row.cells[0].text.strip())  # 第一列数据（姓名）
            use_purpose = str(row.cells[1].text.strip())  # 第二列数据（用途）
            money_amount = str(row.cells[2].text.strip())  # 第三列数据（金额）
            if student_name not in student_name_datas:
                student_name_datas.append(student_name)
            use_purpose_idx = use_purpose_order_switch(use_purpose)
            detail_word_document_datas.append([student_name, use_purpose, money_amount, row_idx, use_purpose_idx])
        
        # 制作 student_name_idx
        for detail_word_document_data in detail_word_document_datas:
            student_name_idx = student_name_datas.index(detail_word_document_data[0])
            if detail_word_document_data[0] == usr_name:
                student_name_idx = -1
            detail_word_document_data.append(student_name_idx)
        # 至此完成排序所需的三个索引号制作：
        # student_name_idx、use_purpose_idx、row_idx

        sorted_detail_word_document_datas = sorted(
            detail_word_document_datas,
            key=lambda x: (x[5], x[4], x[3]) # student_name_idx、use_purpose_idx、row_idx
        )
        
        # 将排序后的数据写回报销明细 Word 文档的对应位置
        for idx, sorted_data in enumerate(sorted_detail_word_document_datas, start=1):
            row = table.rows[idx]
            row.cells[0].text = sorted_data[0]  # 姓名
            row.cells[1].text = sorted_data[1]  # 用途
            row.cells[2].text = sorted_data[2]  # 金额
            detail_word_document_table_cell_style_set(row.cells[0])
            detail_word_document_table_cell_style_set(row.cells[1])
            detail_word_document_table_cell_style_set(row.cells[2])

        # 保存文档
        doc.save(detail_word_document_path)
    
    return

# 更新报销明细 Word 文档后缀
def update_detail_word_document_suffix(detail_word_document_path):
    # 获取当前日期
    current_date = datetime.now().strftime("%Y%m%d")
    
    # 拆分文件路径
    dir, detail_word_document_name = os.path.split(detail_word_document_path)
    detail_word_document_name_without_ext, ext = os.path.splitext(detail_word_document_name)
    
    # 找到后缀部分并替换
    updated_detail_word_document_name_without_ext = re.sub(r'_\d{8}$', f'_{current_date}', detail_word_document_name_without_ext)
      
    # 构造新的文件路径
    updated_detail_word_document_path = os.path.join(dir, updated_detail_word_document_name_without_ext + ext)

    os.rename(detail_word_document_path, updated_detail_word_document_path)
    
    return updated_detail_word_document_path

# 第二个功能小函数结束
# use_purpose_order_switch(use_purpose)
# sort_detail_word_document(detail_word_document_path, usr_name)
# update_detail_word_document_suffix(detail_word_document_path)

# 第二个功能大函数开始
# 在报销明细 Word 文档中写入数据
def write_detail_word_document(work_space, information_excel_document_path, usr_name, student_name, use_purpose, money_amount):
    
    # 搜索当前月文件夹
    try:
        target_numeric_folder_path = find_target_numeric_folder(work_space)
        print_info_find_target_numeric_folder(work_space, target_numeric_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_target_numeric_folder(e)
        return

    # 搜索报销明细 Word 文档
    try:
        detail_word_document_path = find_detail_word_document(target_numeric_folder_path)
        print_info_find_detail_word_document(target_numeric_folder_path, detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_detail_word_document(e)
        return

    # 写入数据
    try:
        # 打开文档
        doc = Document(detail_word_document_path)
        # 获取文档中的表格列表（检查是否已被初始化）
        tables = doc.tables
        # 保存文档
        doc.save(detail_word_document_path)

        if not tables:
            initialize_detail_word_document(detail_word_document_path, information_excel_document_path, usr_name)
        
        # 再次打开文档
        doc = Document(detail_word_document_path)
        # 获取文档中的表格列表（检查是否已被初始化）
        tables = doc.tables
        if tables:
            # 使用顿号（、）分割字符串 money_amount
            money_amounts = money_amount.split('、')
            for amount in money_amounts:
                detail_word_document_add_table(doc, student_name, use_purpose, amount)
                print_info_detail_word_document_add_table(student_name, use_purpose, amount)
        # 保存文档
        doc.save(detail_word_document_path)

    except Exception as e:
        print_ten_pentagram()
        print_error_detail_word_document_add_table(detail_word_document_path, e)
        return

    # 每次写入后都要重排
    try:
        sort_detail_word_document(detail_word_document_path, usr_name)
        print_info_sort_detail_word_document(detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_sort_detail_word_document(detail_word_document_path, e)
        return

    # 每次写入后都要更新后缀
    try:
        updated_detail_word_document_path = update_detail_word_document_suffix(detail_word_document_path)
        print_info_update_detail_word_document_suffix(detail_word_document_path, updated_detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_update_detail_word_document_suffix(detail_word_document_path, e)
        return

    # 输出：1、确认信息；2、姓名和学号
    try:
        student_id = find_student_id_by_name(information_excel_document_path, student_name)
        print_info_check(student_name, use_purpose, money_amount)
        print_info_student_name_with_id(student_name, student_id)
    except Exception as e:
        print_ten_pentagram()
        print_error_student_name_with_id(student_name, e)
        return
# 第二个功能大函数结束

# 第三个功能小函数开始
# find_detail_excel_document(target_folder_path)
# initialize_detail_excel_document(detail_excel_document_path, padding, sheet_name = "报销明细")
# use_purpose_mapping_switch(use_purpose)
# delete_sum_detail_word_document(detail_word_document_path)
# read_detail_word_document(detail_word_document_path)

# write_detail_excel_document(detail_excel_document_path, detail_word_document_datas, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_personal(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_second_table(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_second_table_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# beautify_detail_excel_document(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")

# find_taxi_explanation_document(target_folder_path)
# taxi_explanation_document_student_name(detail_word_document_datas)
# explanation_document_add_title(doc, text)
# explanation_document_add_paragraph(doc, text)
# explanation_document_add_inscribed(doc, text)
# explanation_document_add_blank(doc)
# write_taxi_explanation_document(taxi_explanation_document_path, detail_word_document_datas, usr_name, destination)
# find_special_material_explanation_document(target_folder_path)
# write_special_material_explanation_document(special_material_explanation_document_path, usr_name, special_material)

# output_documents(target_numeric_folder_path, output_space)

# 找当月报销明细 Excel 文档   yes：删除再创建一个；no：创建一个
def find_detail_excel_document(target_folder_path):
    # 搜索目录
    try:
        ensure_path_exist(target_folder_path)
        print_info_path_found(target_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_path_found(target_folder_path)
        raise FileNotFoundError

    # 搜索目录中的 Excel 文件
    detail_excel_documents = [file for file in os.listdir(target_folder_path) if file.endswith(".xlsx") and "报销明细" in file]
    
    # 若不存在，则创建新的 Excel 文档
    # 应当注意可能在汇总时已经到下一个月，需要提取 target_folder_path 中的年月信息
    target_folder_name = os.path.basename(target_folder_path)
    target_folder_year = target_folder_name[:4]
    target_folder_month = int(target_folder_name[4:])
    target_folder_year_month = f"{target_folder_year}年{target_folder_month}月"
    current_year_month_date_num = datetime.now().strftime("%Y%m%d")
    detail_excel_document = f"{target_folder_year_month}报销明细_{current_year_month_date_num}.xlsx"

    # 如果存在符合条件的文件，先删除
    if detail_excel_documents:
        for file in detail_excel_documents:
            os.remove(os.path.join(target_folder_path, file))

    # 创建并保存空白 Excel 文档
    wb = Workbook()
    ws = wb.active
    ws.title = "报销明细"
    wb.save(os.path.join(target_folder_path, detail_excel_document))

    # 返回报销明细 Excel 文档的绝对路径
    detail_excel_document_path = os.path.join(target_folder_path, detail_excel_document)
    return detail_excel_document_path

# 初始化报销明细 Excel 文档
def initialize_detail_excel_document(detail_excel_document_path, padding, sheet_name = "报销明细"):
    wb = load_workbook(detail_excel_document_path)
    # 选择指定的工作表
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        raise ValueError
    
    # 准备传递参数
    initial_datas_dict = {}

    if padding:
        start_row = 2
        start_col = 2
    else:
        start_row = 1
        start_col = 1
    initial_datas_dict['start_row'] = start_row
    initial_datas_dict['start_col'] = start_col

    ws.cell(row=start_row, column=start_col, value="姓名")
    ws.cell(row=start_row, column=start_col + 1, value="用途")
    ws.cell(row=start_row, column=start_col + 2, value="金额")
    ws.cell(row=start_row, column=start_col + 3, value="类别")
    ws.cell(row=start_row, column=start_col + 4, value="个人汇总")

    # 保存工作簿
    wb.save(detail_excel_document_path)
    return initial_datas_dict

# 用途映射 switch
def use_purpose_mapping_switch(use_purpose):
    use_purpose_mapping_dict = {
        '出租车费': '市内交通',
        '过路费': '市内交通',
        '过桥费': '市内交通',
        '地铁费': '市内交通',
        '寄件运费': '物流',
        '飞机票': '差旅',
        '高铁票': '差旅',
        '住宿费': '差旅',
        '餐饮费': '差旅',
        '论文费': '论文',
        '打印费': '打印',
    }
    return use_purpose_mapping_dict.get(use_purpose, '材料')

# 报销明细 Word 文档删除合计行（防止重复操作 bug）
def delete_sum_detail_word_document(detail_word_document_path):
    # 打开Word文档
    doc = Document(detail_word_document_path)
    # 获取文档中的表格列表
    tables = doc.tables
    if tables:
        table = tables[0]   
        
        # 找到并删除合计行
        rows_to_delete = []
        for row_idx, row in enumerate(table.rows[1:], start=1):  # 从第二行开始遍历
            student_name = str(row.cells[0].text.strip())  # 第一列数据（姓名）
            if student_name == '合计':
                rows_to_delete.append(row_idx)  # 记录需要删除的行索引

        # 删除记录中的行
        for row_idx in reversed(rows_to_delete):  # 从后往前删除，避免影响索引
            table._element.remove(table.rows[row_idx]._element)
        
        # 保存修改后的文档
        doc.save(detail_word_document_path)
    return

# 读取报销明细 Word 文档
def read_detail_word_document(detail_word_document_path):
    # 打开Word文档
    doc = Document(detail_word_document_path)
    # 获取文档中的表格列表
    tables = doc.tables
    if tables:
        table = tables[0]
        # 初始化数组
        detail_word_document_datas = []
        # 遍历表格的每一行，从第二行开始（跳过表头）
        for row_idx, row in enumerate(table.rows[1:], start=2):
            # 获取当前行的三个单元格数据
            student_name = str(row.cells[0].text.strip())  # 第一列数据（姓名）
            if student_name == '合计':
                pass
                continue
            use_purpose = str(row.cells[1].text.strip())  # 第二列数据（用途）
            money_amount = str(row.cells[2].text.strip())  # 第三列数据（金额）
            detail_word_document_datas.append([student_name, use_purpose, money_amount, use_purpose_mapping_switch(use_purpose)])
    return detail_word_document_datas

# 写入报销明细 Excel 文档 1/6
def write_detail_excel_document(detail_excel_document_path, detail_word_document_datas, initial_datas_dict, sheet_name = "报销明细"):
    wb = load_workbook(detail_excel_document_path)
    # 选择指定的工作表
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        raise ValueError

    # start_row - 起始行
    # start_col - 起始列
    # detail_word_document_datas_length - 数据总行数（大表格数据行数）
    # use_purpose_mapping_datas_length - 用途数（小表格数据行数）
    # use_purpose_mapping_datas - 用途列表
    # initial_sum - 原始合计
    # student_number - 学生数
    # student_names - 学生列表
    # sum_row - 大表格合计行
    # second_table_row - 小表格起始行
    # second_table_sum_row - 小表格合计行

    # 传递参数：大表格数据行数
    detail_word_document_datas_length = len(detail_word_document_datas)
    initial_datas_dict['detail_word_document_datas_length'] = detail_word_document_datas_length

    # 传递参数：小表格数据
    use_purpose_mapping_datas = []

    # 传递参数：原始总计
    initial_sum = 0

    # 从第二行开始写入数据
    for idx, data in enumerate(detail_word_document_datas, start=initial_datas_dict['start_row'] + 1):
        ws.cell(row=idx, column=initial_datas_dict['start_col'], value=data[0])  # 写入姓名
        ws.cell(row=idx, column=initial_datas_dict['start_col'] + 1, value=data[1])  # 写入用途

        # 写入金额
        if isinstance(data[2], (int, float)):
            data_2 = data[2]
        else:
            data_2 = float(data[2])
        ws.cell(row=idx, column=initial_datas_dict['start_col'] + 2, value=data_2)
        initial_sum = initial_sum + data_2

        ws.cell(row=idx, column=initial_datas_dict['start_col'] + 3, value=data[3])  # 写入类别
        if data[3] not in use_purpose_mapping_datas:
            use_purpose_mapping_datas.append(data[3])

    use_purpose_mapping_datas_length = len(use_purpose_mapping_datas)
    initial_datas_dict['use_purpose_mapping_datas_length'] = use_purpose_mapping_datas_length
    initial_datas_dict['use_purpose_mapping_datas'] = use_purpose_mapping_datas
    initial_sum = round(initial_sum, 2)
    initial_datas_dict['initial_sum'] = initial_sum

    # 合并姓名、个人汇总列单元格
    student_name_column = initial_datas_dict['start_col']
    personal_sum_column = initial_datas_dict['start_col'] + 4

    # 获取姓名列中所有单元格对象
    student_name_cells = list(ws.iter_cols(
        min_col=student_name_column,
        max_col=student_name_column,
        min_row=initial_datas_dict['start_row'] + 1,
        max_row=detail_word_document_datas_length + initial_datas_dict['start_row']
    ))[0]

    # 遍历列并合并重复的单元格
    start_row = initial_datas_dict['start_row'] + 1
    student_names = []  # 传递参数：人名
    student_names.append(student_name_cells[0].value)
    student_number = 1  # 传递参数：人数

    for i in range(1, len(student_name_cells)):
        current_cell = student_name_cells[i]  # 当前单元格
        previous_cell = student_name_cells[i - 1]  # 前一个单元格
        current_cell_value = current_cell.value
        previous_cell_value = previous_cell.value

        # 如果当前单元格的值与前一个单元格相同
        if current_cell_value == previous_cell_value:
            if i == (len(student_name_cells) - 1):
                ws.merge_cells(
                    start_row=start_row,
                    start_column=student_name_column,
                    end_row=i + initial_datas_dict['start_row'] + 1,
                    end_column=student_name_column
                )
                ws.merge_cells(
                    start_row=start_row,
                    start_column=personal_sum_column,
                    end_row=i + initial_datas_dict['start_row'] + 1,
                    end_column=personal_sum_column
                )
        else:
            student_number += 1
            student_names.append(current_cell_value)
            ws.merge_cells(
                start_row=start_row,
                start_column=student_name_column,
                end_row=i + initial_datas_dict['start_row'],
                end_column=student_name_column
            )
            ws.merge_cells(
                start_row=start_row,
                start_column=personal_sum_column,
                end_row=i + initial_datas_dict['start_row'],
                end_column=personal_sum_column
            )
            start_row = i + initial_datas_dict['start_row'] + 1

    initial_datas_dict['student_number'] = student_number
    initial_datas_dict['student_names'] = student_names

    sum_row = detail_word_document_datas_length + initial_datas_dict['start_row'] + 1
    initial_datas_dict['sum_row'] = sum_row

    second_table_row = sum_row + 2
    initial_datas_dict['second_table_row'] = second_table_row
    second_table_sum_row = second_table_row + 1 + use_purpose_mapping_datas_length
    initial_datas_dict['second_table_sum_row'] = second_table_sum_row

    ws.cell(row=sum_row, column=initial_datas_dict['start_col'], value="合计")

    # 小表格
    ws.cell(row=second_table_row, column=initial_datas_dict['start_col'], value="分类号")
    ws.cell(row=second_table_row, column=initial_datas_dict['start_col'] + 1, value="类别")
    ws.cell(row=second_table_row, column=initial_datas_dict['start_col'] + 2, value="小计")
    for idx, data in enumerate(use_purpose_mapping_datas, start=1):
        ws.cell(row=second_table_row+idx, column=initial_datas_dict['start_col'], value=idx)  # 写入分类号
        ws.cell(row=second_table_row+idx, column=initial_datas_dict['start_col'] + 1, value=data)  # 写入类别
    ws.cell(row=second_table_sum_row, column=initial_datas_dict['start_col'], value="合计")

    wb.save(detail_excel_document_path)
    return initial_datas_dict

# 计算报销明细 Excel 文档：个人汇总 2/6
def calculate_detail_excel_document_personal(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细"):
    wb = load_workbook(detail_excel_document_path)
    # 选择指定的工作表
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        raise ValueError

    # 准备传递数据：student_sum
    student_sums = []

    row = initial_datas_dict['start_row'] + 1
    # 处理 student_number 次就结束
    for i in range(initial_datas_dict['student_number']):
        
        # 判断该格是否是合并单元格
        merged = False
        for merged_cell in ws.merged_cells.ranges:
            if ws.cell(row=row, column=initial_datas_dict['start_col'] + 4).coordinate in merged_cell:
                merged = True
                break
        
        if merged:
            # 如果是合并单元格，计算该合并单元格范围的总和
            min_row = merged_cell.min_row
            max_row = merged_cell.max_row
            sum_value = 0
            for r in range(min_row, max_row + 1):
                value = ws.cell(row=r, column=initial_datas_dict['start_col'] + 2).value

                # 确保 value 是数字类型
                if isinstance(value, (int, float)):
                    sum_value += value
                else:
                    sum_value += float(value)  # 尝试转换为 float 类型
            # 填入总和
            sum_value = round(sum_value, 2)
            ws.cell(row=min_row, column=initial_datas_dict['start_col'] + 4, value=sum_value)
            student_sums.append(sum_value)
            row = max_row + 1
        
        else:
            # 如果不是合并单元格，直接填入金额列的值
            sum_value = ws.cell(row=row, column=initial_datas_dict['start_col'] + 2).value

            if isinstance(sum_value, (int, float)):
                ws.cell(row=row, column=initial_datas_dict['start_col'] + 4, value=sum_value)
            else:
                ws.cell(row=row, column=initial_datas_dict['start_col'] + 4, value=float(sum_value))
            sum_value = round(sum_value, 2)
            student_sums.append(sum_value)
            row = row + 1

    initial_datas_dict['student_sums'] = student_sums
    # 保存修改后的文件
    wb.save(detail_excel_document_path)
    return initial_datas_dict

# start_row - 起始行
# start_col - 起始列
# detail_word_document_datas_length - 数据总行数（大表格数据行数）
# use_purpose_mapping_datas_length - 用途数（小表格数据行数）
# use_purpose_mapping_datas - 用途列表
# initial_sum - 原始合计
# student_number - 学生数
# student_names - 学生列表
# sum_row - 大表格合计行
# second_table_row - 小表格起始行
# second_table_sum_row - 小表格合计行
# student_sums - 个人汇总

# 计算报销明细 Excel 文档：大表格合计 3/6
def calculate_detail_excel_document_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细"):
    wb = load_workbook(detail_excel_document_path)
    # 选择指定的工作表
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        raise ValueError
    
    initial_sum = initial_datas_dict['initial_sum']
    student_sum = sum(initial_datas_dict['student_sums'])
    student_sum = round(student_sum, 2)

    if abs(initial_sum - student_sum) >= 0.001:
        raise ValueError()

    ws.cell(row=initial_datas_dict['sum_row'], column=initial_datas_dict['start_col'] + 2, value=initial_sum)
    ws.cell(row=initial_datas_dict['sum_row'], column=initial_datas_dict['start_col'] + 4, value=student_sum)

    # 保存修改后的文件
    wb.save(detail_excel_document_path)
    return initial_datas_dict

# start_row - 起始行
# start_col - 起始列
# detail_word_document_datas_length - 数据总行数（大表格数据行数）
# use_purpose_mapping_datas_length - 用途数（小表格数据行数）
# use_purpose_mapping_datas - 用途列表
# initial_sum - 原始合计
# student_number - 学生数
# student_names - 学生列表
# sum_row - 大表格合计行
# second_table_row - 小表格起始行
# second_table_sum_row - 小表格合计行
# student_sums - 个人汇总
# purpose_sums - 用途汇总

# 计算报销明细 Excel 文档：小表格小计 4/6
def calculate_detail_excel_document_second_table(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细"):
    wb = load_workbook(detail_excel_document_path)
    # 选择指定的工作表
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        raise ValueError

    use_purpose_mapping_datas = initial_datas_dict['use_purpose_mapping_datas']
    purpose_sums = {}

    for use_purpose_mapping_data in use_purpose_mapping_datas:
        # 初始化
        purpose_sums[use_purpose_mapping_data] = 0
        for row in range(initial_datas_dict['start_row'] + 1, initial_datas_dict['sum_row']):
            if ws.cell(row=row, column=initial_datas_dict['start_col'] + 3).value == use_purpose_mapping_data:
                value_of_use_purpose_mapping_data = ws.cell(row=row, column=initial_datas_dict['start_col'] + 2).value
                if isinstance(value_of_use_purpose_mapping_data, (int, float)):
                    purpose_sums[use_purpose_mapping_data] += value_of_use_purpose_mapping_data
                else:
                    purpose_sums[use_purpose_mapping_data] += float(value_of_use_purpose_mapping_data)
                purpose_sums[use_purpose_mapping_data] = round(purpose_sums[use_purpose_mapping_data], 2)

    initial_datas_dict['purpose_sums'] = purpose_sums

    for row in range(initial_datas_dict['second_table_row'] + 1, initial_datas_dict['second_table_sum_row']):
        use_purpose_mapping_data = ws.cell(row=row, column=initial_datas_dict['start_col'] + 1).value
        purpose_sum = purpose_sums[use_purpose_mapping_data]
        if isinstance(purpose_sum, (int, float)):
            ws.cell(row=row, column=initial_datas_dict['start_col'] + 2, value=purpose_sum)
        else:
            ws.cell(row=row, column=initial_datas_dict['start_col'] + 2, value=float(purpose_sum))

    # 保存修改后的文件
    wb.save(detail_excel_document_path)
    return initial_datas_dict

# start_row - 起始行
# start_col - 起始列
# detail_word_document_datas_length - 数据总行数（大表格数据行数）
# use_purpose_mapping_datas_length - 用途数（小表格数据行数）
# use_purpose_mapping_datas - 用途列表
# initial_sum - 原始合计
# student_number - 学生数
# student_names - 学生列表
# sum_row - 大表格合计行
# second_table_row - 小表格起始行
# second_table_sum_row - 小表格合计行
# student_sums - 个人汇总
# purpose_sums - 用途汇总

# 计算报销明细 Excel 文档：小表格合计 5/6
def calculate_detail_excel_document_second_table_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细"):
    wb = load_workbook(detail_excel_document_path)
    # 选择指定的工作表
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        raise ValueError

    initial_sum = initial_datas_dict['initial_sum']
    purpose_sums = initial_datas_dict['purpose_sums']
    purpose_sum = sum(purpose_sums.values())
    purpose_sum = round(purpose_sum, 2)

    if abs(initial_sum - purpose_sum) >= 0.001:
        raise ValueError

    ws.cell(row=initial_datas_dict['second_table_sum_row'], column=initial_datas_dict['start_col'] + 2, value=initial_sum)

    # 保存修改后的文件
    wb.save(detail_excel_document_path)
    return initial_datas_dict

# start_row - 起始行
# start_col - 起始列
# detail_word_document_datas_length - 数据总行数（大表格数据行数）
# use_purpose_mapping_datas_length - 用途数（小表格数据行数）
# use_purpose_mapping_datas - 用途列表
# initial_sum - 原始合计
# student_number - 学生数
# student_names - 学生列表
# sum_row - 大表格合计行
# second_table_row - 小表格起始行
# second_table_sum_row - 小表格合计行
# student_sums - 个人汇总
# purpose_sums - 用途汇总

# 渲染报销明细 Excel 文档 6/6
def beautify_detail_excel_document(detail_excel_document_path, initial_datas_dict, colors, sheet_name = "报销明细"):
    wb = load_workbook(detail_excel_document_path)
    # 选择指定的工作表
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        raise ValueError

    # 设置边框样式：0.5磅
    border_style = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thin", color="000000"),
        bottom=Side(border_style="thin", color="000000")
    )

    border_style_left_thick = Border(
        left=Side(border_style="thick", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thin", color="000000"),
        bottom=Side(border_style="thin", color="000000")
    )

    border_style_right_thick = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thick", color="000000"),
        top=Side(border_style="thin", color="000000"),
        bottom=Side(border_style="thin", color="000000")
    )

    border_style_bottom_thick = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thin", color="000000"),
        bottom=Side(border_style="thick", color="000000")
    )

    border_style_right_top_bottom_thick = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thick", color="000000"),
        top=Side(border_style="thick", color="000000"),
        bottom=Side(border_style="thick", color="000000")
    )

    border_style_top_bottom_thick = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thick", color="000000"),
        bottom=Side(border_style="thick", color="000000")
    )

    border_style_left_top_bottom_thick = Border(
        left=Side(border_style="thick", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thick", color="000000"),
        bottom=Side(border_style="thick", color="000000")
    )

    border_style_left_top_thick = Border(
        left=Side(border_style="thick", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thick", color="000000"),
        bottom=Side(border_style="thin", color="000000")
    )

    border_style_left_bottom_thick = Border(
        left=Side(border_style="thick", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thin", color="000000"),
        bottom=Side(border_style="thick", color="000000")
    )

    border_style_right_top_thick = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thick", color="000000"),
        top=Side(border_style="thick", color="000000"),
        bottom=Side(border_style="thin", color="000000")
    )

    border_style_right_bottom_thick = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thick", color="000000"),
        top=Side(border_style="thin", color="000000"),
        bottom=Side(border_style="thick", color="000000")
    )

    # 设置对齐样式：水平和垂直居中
    alignment_style = Alignment(horizontal='center', vertical='center')

    # 设置字体样式
    font_style_title = Font(name='等线', size=12, bold=True)
    font_style_sum = Font(name='等线', size=11, bold=True)
    font_style = Font(name='等线', size=11, bold=False)

    colors = [color.lstrip('#') for color in colors]

    fill_material = PatternFill(start_color=colors[0], end_color=colors[0], fill_type="solid") # 材料
    fill_transportation = PatternFill(start_color=colors[1], end_color=colors[1], fill_type="solid") # 市内交通
    fill_logistics = PatternFill(start_color=colors[2], end_color=colors[2], fill_type="solid") # 物流
    fill_print = PatternFill(start_color=colors[3], end_color=colors[3], fill_type="solid") # 打印
    fill_travel = PatternFill(start_color=colors[4], end_color=colors[4], fill_type="solid") # 差旅
    fill_paper = PatternFill(start_color=colors[5], end_color=colors[5], fill_type="solid") # 论文
    
    for row in range(initial_datas_dict['start_row'], initial_datas_dict['sum_row'] + 1):
        for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 5):
            cell = ws.cell(row=row, column=col)
            cell.alignment = alignment_style
            cell.border = border_style
            cell.font = font_style

    for row in range(initial_datas_dict['second_table_row'], initial_datas_dict['second_table_sum_row'] + 1):
        for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 3):
            cell = ws.cell(row=row, column=col)
            cell.alignment = alignment_style
            cell.border = border_style
            cell.font = font_style

    row = initial_datas_dict['start_row']
    for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 5):
        cell = ws.cell(row=row, column=col)
        cell.font = font_style_title

    row = initial_datas_dict['second_table_row']
    for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 3):
        cell = ws.cell(row=row, column=col)
        cell.font = font_style_title

    row = initial_datas_dict['sum_row']
    for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 5):
        cell = ws.cell(row=row, column=col)
        cell.font = font_style_sum

    row = initial_datas_dict['second_table_sum_row']
    for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 3):
        cell = ws.cell(row=row, column=col)
        cell.font = font_style_sum

    if initial_datas_dict['start_col'] == 1:
        # 设置列宽
        column_widths = {
        'A': 12,
        'B': 20,
        'C': 14,
        'D': 14,
        'E': 14
        }
        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

    if initial_datas_dict['start_col'] == 2:
        # 设置列宽
        column_widths = {
        'A': 3,
        'B': 12,
        'C': 20,
        'D': 14,
        'E': 14,
        'F': 14,
        }
        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

    for row in range(1, initial_datas_dict['second_table_sum_row'] + 1):
        ws.row_dimensions[row].height = 16

    for row in range(initial_datas_dict['start_row'] + 2, initial_datas_dict['sum_row']):
        previous_student_name = ws.cell(row=row - 1, column=initial_datas_dict['start_col']).value
        current_student_name = ws.cell(row=row, column=initial_datas_dict['start_col']).value
        if previous_student_name and (not current_student_name):
            continue
        if current_student_name != previous_student_name:
            for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 5):
                cell = ws.cell(row=row - 1, column=col)
                cell.border = border_style_bottom_thick
    
    for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 5):
        cell = ws.cell(row=initial_datas_dict['start_row'], column=col)
        cell.border = border_style_bottom_thick
        cell = ws.cell(row=initial_datas_dict['sum_row'] - 1, column=col)
        cell.border = border_style_bottom_thick
        cell = ws.cell(row=initial_datas_dict['sum_row'], column=col)
        cell.border = border_style_bottom_thick

    for row in range(initial_datas_dict['start_row'], initial_datas_dict['sum_row'] + 1):
        cell = ws.cell(row=row, column=initial_datas_dict['start_col'] + 4)
        cell.border = border_style_right_top_bottom_thick
    
    for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 5):
        cell = ws.cell(row=initial_datas_dict['start_row'], column=col)
        cell.border = border_style_top_bottom_thick
    
    for row in range(initial_datas_dict['start_row'], initial_datas_dict['sum_row'] + 1):
        cell = ws.cell(row=row, column=initial_datas_dict['start_col'])
        cell.border = border_style_left_top_bottom_thick

    ws.cell(row=initial_datas_dict['start_row'], column=initial_datas_dict['start_col'] + 4).border = border_style_right_top_bottom_thick

    # 为小表格的每个边设置边框
    for col in range(initial_datas_dict['start_col'], initial_datas_dict['start_col'] + 3):
        cell = ws.cell(row=initial_datas_dict['second_table_row'], column=col)
        cell.border = border_style_top_bottom_thick
        cell = ws.cell(row=initial_datas_dict['second_table_sum_row'], column=col)
        cell.border = border_style_top_bottom_thick

    for row in range(initial_datas_dict['second_table_row'], initial_datas_dict['second_table_sum_row'] + 1):
        cell = ws.cell(row=row, column=initial_datas_dict['start_col'])
        if row == initial_datas_dict['second_table_row'] or row == initial_datas_dict['second_table_sum_row']:
            cell.border = border_style_left_top_bottom_thick
        elif row == initial_datas_dict['second_table_row'] + 1:
            cell.border = border_style_left_top_thick
        elif row == initial_datas_dict['second_table_sum_row'] - 1:
            cell.border = border_style_left_bottom_thick
        else:
            cell.border = border_style_left_thick
        cell = ws.cell(row=row, column=initial_datas_dict['start_col'] + 2)
        if row == initial_datas_dict['second_table_row'] or row == initial_datas_dict['second_table_sum_row']:
            cell.border = border_style_right_top_bottom_thick
        elif row == initial_datas_dict['second_table_row'] + 1:
            cell.border = border_style_right_top_thick
        elif row == initial_datas_dict['second_table_sum_row'] - 1:
            cell.border = border_style_right_bottom_thick
        else:
            cell.border = border_style_right_thick

    # 改变大表格用途列背景颜色
    col = initial_datas_dict['start_col'] + 3
    for row in range(initial_datas_dict['start_row'] + 1, initial_datas_dict['sum_row']):
        cell = ws.cell(row=row, column=col)
        if cell.value == '材料':
            cell.fill = fill_material
        elif cell.value == '市内交通':
            cell.fill = fill_transportation
        elif cell.value == '物流':
            cell.fill = fill_logistics
        elif cell.value == '打印':
            cell.fill = fill_print
        elif cell.value == '差旅':
            cell.fill = fill_travel
        elif cell.value == '论文':
            cell.fill = fill_paper
    
    # 改变小表格用途列背景颜色
    col = initial_datas_dict['start_col'] + 1
    for row in range(initial_datas_dict['second_table_row'] + 1, initial_datas_dict['second_table_sum_row']):
        cell = ws.cell(row=row, column=col)
        if cell.value == '材料':
            cell.fill = fill_material
        elif cell.value == '市内交通':
            cell.fill = fill_transportation
        elif cell.value == '物流':
            cell.fill = fill_logistics
        elif cell.value == '打印':
            cell.fill = fill_print
        elif cell.value == '差旅':
            cell.fill = fill_travel
        elif cell.value == '论文':
            cell.fill = fill_paper

    # 保存修改后的文件
    wb.save(detail_excel_document_path)
    return initial_datas_dict

# 搜索出租车情况说明文档
def find_taxi_explanation_document(target_folder_path):
    # 搜索目录
    try:
        ensure_path_exist(target_folder_path)
        print_info_path_found(target_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_path_found(target_folder_path)
        raise FileNotFoundError

    taxi_explanation_documents = [file for file in os.listdir(target_folder_path) if file.endswith(".docx") and "出租车情况说明" in file]
    
    # 若不存在，则创建新的 Word 文档
    # 应当注意可能在汇总时已经到下一个月，需要提取 target_folder_path 中的年月信息
    target_folder_name = os.path.basename(target_folder_path)
    target_folder_year = target_folder_name[:4]
    target_folder_month = int(target_folder_name[4:])
    target_folder_year_month = f"{target_folder_year}年{target_folder_month}月"
    current_year_month_date_num = datetime.now().strftime("%Y%m%d")
    taxi_explanation_document = f"{target_folder_year_month}出租车情况说明_{current_year_month_date_num}.docx"

    # 如果存在符合条件的文件，先删除
    if taxi_explanation_documents:
        for file in taxi_explanation_documents:
            os.remove(os.path.join(target_folder_path, file))

    # 创建并保存空白文档
    doc = Document()
    doc.save(os.path.join(target_folder_path, taxi_explanation_document))

    # 返回报销明细 Word 文档的绝对路径
    taxi_explanation_document_path = os.path.join(target_folder_path, taxi_explanation_document)
    return taxi_explanation_document_path

# 获取出租车学生姓名
def taxi_explanation_document_student_name(detail_word_document_datas):
    student_names = []
    for data in detail_word_document_datas:
        if (data[3] == '市内交通') and (data[0] not in student_names):
            student_names.append(data[0])
    return student_names

# 写入情况说明文档标题
def explanation_document_add_title(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph()
    # 水平居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 设置段前/段后间距
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1  # 单倍行距
    paragraph_format.space_before = Pt(30)
    paragraph_format.space_after = Pt(30)

    # 设置字号
    run = paragraph.add_run(text)
    run.font.size = Pt(30)

    # 设置中文为宋体，英文及数字为新罗马
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.append(rFonts)

# 写入情况说明文档段落
def explanation_document_add_paragraph(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph()
    # 水平居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # 设置段前/段后间距
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    # 设置字号
    run = paragraph.add_run(text)
    run.font.size = Pt(20)

    # 设置中文为宋体，英文及数字为新罗马
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.append(rFonts)

# 写入情况说明文档落款
def explanation_document_add_inscribed(doc, text):
    # 插入段落
    paragraph = doc.add_paragraph()
    # 水平居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # 设置段前/段后间距
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    # 设置字号
    run = paragraph.add_run(text)
    run.font.size = Pt(20)

    # 设置中文为宋体，英文及数字为新罗马
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.append(rFonts)

# 写入情况说明文档空行
def explanation_document_add_blank(doc):
    # 插入段落
    paragraph = doc.add_paragraph()
    # 水平居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # 设置段前/段后间距
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    # 设置字号
    run = paragraph.add_run("")
    run.font.size = Pt(10)

    # 设置中文为宋体，英文及数字为新罗马
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.append(rFonts)

# 写入出租车情况说明文档
def write_taxi_explanation_document(taxi_explanation_document_path, detail_word_document_datas, usr_name, destination):
    
    doc = Document(taxi_explanation_document_path)

    # 获取出租车学生姓名
    student_names = taxi_explanation_document_student_name(detail_word_document_datas)

    student_names_str = '、'.join(student_names)

    # 写入出租车情况说明文档标题
    explanation_document_add_title(doc, "出租车情况说明")
    print_info_explanation_document_add_title('出租车', taxi_explanation_document_path)

    # 写入出租车情况说明文档段落
    explanation_document_add_paragraph(doc, "　　" + student_names_str + "等人需要往返于实验室与" + destination + "等地，进行项目汇报、项目评审、设备调试、机器修理、学术交流、人员接洽等工作。")
    explanation_document_add_blank(doc)
    explanation_document_add_blank(doc)
    explanation_document_add_paragraph(doc, "　　特此说明")
    explanation_document_add_blank(doc)
    explanation_document_add_blank(doc)
    print_info_explanation_document_add_paragraph('出租车', taxi_explanation_document_path)

    # 写入出租车情况说明文档落款
    explanation_document_add_inscribed(doc, "负责人：" + usr_name)
    explanation_document_add_inscribed(doc, datetime.now().strftime("%Y年") + str(int(datetime.now().strftime("%m"))) + "月" + str(int(datetime.now().strftime("%d"))) + "日")
    print_info_explanation_document_add_inscribed('出租车', taxi_explanation_document_path)

    # 保存文档
    doc.save(taxi_explanation_document_path)
    return

# 搜索不建固按材料报销情况说明文档
def find_special_material_explanation_document(target_folder_path):
    # 搜索目录
    try:
        ensure_path_exist(target_folder_path)
        print_info_path_found(target_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_path_found(target_folder_path)
        raise FileNotFoundError

    special_material_explanation_documents = [file for file in os.listdir(target_folder_path) if file.endswith(".docx") and "不建固按材料报销情况说明" in file]
    
    # 若不存在，则创建新的 Word 文档
    # 应当注意可能在汇总时已经到下一个月，需要提取 target_folder_path 中的年月信息
    target_folder_name = os.path.basename(target_folder_path)
    target_folder_year = target_folder_name[:4]
    target_folder_month = int(target_folder_name[4:])
    target_folder_year_month = f"{target_folder_year}年{target_folder_month}月"
    current_year_month_date_num = datetime.now().strftime("%Y%m%d")
    special_material_explanation_document = f"{target_folder_year_month}不建固按材料报销情况说明_{current_year_month_date_num}.docx"

    # 如果存在符合条件的文件，先删除
    if special_material_explanation_documents:
        for file in special_material_explanation_documents:
            os.remove(os.path.join(target_folder_path, file))

    # 创建并保存空白文档
    doc = Document()
    doc.save(os.path.join(target_folder_path, special_material_explanation_document))

    # 返回报销明细 Word 文档的绝对路径
    special_material_explanation_document_path = os.path.join(target_folder_path, special_material_explanation_document)
    return special_material_explanation_document_path

# 写入不建固按材料报销情况说明文档
def write_special_material_explanation_document(special_material_explanation_document_path, usr_name, special_materials, special_material_uses):

    doc = Document(special_material_explanation_document_path)

    # 定义英文标点和对应的中文标点
    english_punctuation = r",.!?;:()[]"
    chinese_punctuation = r"，。！？；：（）【】"

    # 使用 maketrans 创建英文标点到中文标点的映射
    punctuation_map = str.maketrans(english_punctuation, chinese_punctuation)

    # 去除头部空格、尾部标点，转换中文标点
    for i, (special_material, special_material_use) in enumerate(zip(special_materials, special_material_uses)):
        if special_material and special_material_use:
            # 先将小数点替换为占位符
            special_material = re.sub(r'(\d)\.(\d)', r'\1<decimal>\2', special_material)
            special_material_use = re.sub(r'(\d)\.(\d)', r'\1<decimal>\2', special_material_use)

            # 进行英文标点到中文标点的转换
            special_material = special_material.translate(punctuation_map)
            special_material_use = special_material_use.translate(punctuation_map)

            # 恢复小数点
            special_material = special_material.replace('<decimal>', '.')
            special_material_use = special_material_use.replace('<decimal>', '.')

            # 去除尾部的中文标点
            special_material = re.sub(r'[，。]+$', '', special_material.strip())
            special_material_use = re.sub(r'[，。]+$', '', special_material_use.strip())

            # 更新到原始列表
            special_materials[i] = special_material
            special_material_uses[i] = special_material_use
    # 写入不建固按材料报销情况说明文档标题
    explanation_document_add_title(doc, "不建固按材料报销情况说明")
    print_info_explanation_document_add_title('不建固按材料报销', special_material_explanation_document_path)

    # 写入不建固按材料报销情况说明文档段落
    for special_material, special_material_use in zip(special_materials, special_material_uses):
        if (special_material is not None) and (special_material_use is not None):
            explanation_document_add_paragraph(doc, f"　　【{special_material}】作为机器人组装所需部件，{special_material_use}，不具备独立设备运行功能，申请免建资产，按材料费报销。")

    explanation_document_add_blank(doc)
    explanation_document_add_blank(doc)
    explanation_document_add_paragraph(doc, "　　特此说明")
    explanation_document_add_blank(doc)
    explanation_document_add_blank(doc)
    print_info_explanation_document_add_paragraph('不建固按材料报销', special_material_explanation_document_path)

    # 写入不建固按材料报销情况说明文档落款
    explanation_document_add_inscribed(doc, "负责人：" + usr_name)
    explanation_document_add_inscribed(doc, datetime.now().strftime("%Y年") + str(int(datetime.now().strftime("%m"))) + "月" + str(int(datetime.now().strftime("%d"))) + "日")
    print_info_explanation_document_add_inscribed('不建固按材料报销', special_material_explanation_document_path)

    # 保存文档
    doc.save(special_material_explanation_document_path)
    return

# 输出全部文档
def output_documents(target_numeric_folder_path, output_space):

    if not os.path.exists(target_numeric_folder_path):
        raise FileNotFoundError
    
    ensure_path_exist(output_space)
    
    # 应当注意可能在汇总时已经到下一个月，需要提取 target_numeric_folder_path 中的年月信息
    target_folder_name = os.path.basename(target_numeric_folder_path)
    target_folder_year = target_folder_name[:4]
    target_folder_month = int(target_folder_name[4:])
    zip_folder_name = f"{target_folder_year}年{target_folder_month}月报销"
    
    # 创建压缩文件路径
    zip_folder_path = os.path.join(output_space, f"{zip_folder_name}.zip")
    
    # 创建压缩文件并添加源目录中的所有文件
    with zipfile.ZipFile(zip_folder_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(target_numeric_folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                # 将文件写入压缩包，保持相对路径
                arcname = os.path.relpath(file_path, start=target_numeric_folder_path)
                zipf.write(file_path, arcname)
    return zip_folder_path

# 第三个功能小函数结束
# find_detail_excel_document(target_folder_path)
# initialize_detail_excel_document(detail_excel_document_path, padding, sheet_name = "报销明细")
# use_purpose_mapping_switch(use_purpose)
# delete_sum_detail_word_document(detail_word_document_path)
# read_detail_word_document(detail_word_document_path)

# write_detail_excel_document(detail_excel_document_path, detail_word_document_datas, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_personal(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_second_table(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# calculate_detail_excel_document_second_table_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
# beautify_detail_excel_document(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")

# find_taxi_explanation_document(target_folder_path)
# taxi_explanation_document_student_name(detail_word_document_datas)
# explanation_document_add_title(doc, text)
# explanation_document_add_paragraph(doc, text)
# explanation_document_add_inscribed(doc, text)
# explanation_document_add_blank(doc)
# write_taxi_explanation_document(taxi_explanation_document_path, detail_word_document_datas, usr_name, destination)
# find_special_material_explanation_document(target_folder_path)
# write_special_material_explanation_document(special_material_explanation_document_path, usr_name, special_material)

# output_documents(target_numeric_folder_path, output_space)

# start_row - 起始行
# start_col - 起始列
# detail_word_document_datas_length - 数据总行数（大表格数据行数）
# use_purpose_mapping_datas_length - 用途数（小表格数据行数）
# use_purpose_mapping_datas - 用途列表
# initial_sum - 原始合计
# student_number - 学生数
# student_names - 学生列表
# sum_row - 大表格合计行
# second_table_row - 小表格起始行
# second_table_sum_row - 小表格合计行
# student_sums - 个人汇总
# purpose_sums - 用途汇总

# 第三个功能大函数开始
# 生成并输出全部文档
def build_and_output_documents(usr_name, work_space, output_space, destination, special_materials, special_material_uses, padding, colors):

    # 搜索当前月文件夹
    try:
        target_numeric_folder_path = find_target_numeric_folder(work_space)
        print_info_find_target_numeric_folder(work_space, target_numeric_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_target_numeric_folder(e)
        return

    # 搜索报销明细 Word 文档
    try:
        detail_word_document_path = find_detail_word_document(target_numeric_folder_path)
        print_info_find_detail_word_document(target_numeric_folder_path, detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_detail_word_document(e)
        return

    # 搜索报销明细 Excel 文档
    try:
        detail_excel_document_path = find_detail_excel_document(target_numeric_folder_path)
        print_info_find_detail_excel_document(target_numeric_folder_path, detail_excel_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_detail_excel_document(e)
        return

    # 初始化报销明细 Excel 文档
    try:
        initial_datas_dict = initialize_detail_excel_document(detail_excel_document_path, padding)
        print_info_initialize_detail_excel_document(detail_excel_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_initialize_detail_excel_document(detail_excel_document_path, e)
        return
    
    # 报销明细 Word 文档删除合计行
    try:
        delete_sum_detail_word_document(detail_word_document_path)
        print_info_delete_sum_detail_word_document(detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_delete_sum_detail_word_document(detail_word_document_path, e)
        return
    
    # 读取报销明细 Word 文档数据
    try:
        detail_word_document_datas = read_detail_word_document(detail_word_document_path)
        print_info_read_detail_word_document(detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_read_detail_word_document(detail_word_document_path, e)
        return

    if detail_word_document_datas:
        # 写入报销明细 Excel 文档 1/6
        try:
            initial_datas_dict = write_detail_excel_document(detail_excel_document_path, detail_word_document_datas, initial_datas_dict)
            print_info_write_detail_excel_document(detail_excel_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_write_detail_excel_document(detail_excel_document_path, e)
            return

        # 计算报销明细 Excel 文档：个人汇总 2/6
        try:
            initial_datas_dict = calculate_detail_excel_document_personal(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
            print_info_calculate_detail_excel_document_personal(detail_excel_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_calculate_detail_excel_document(detail_excel_document_path, e)
            return

        # 计算报销明细 Excel 文档：大表格合计 3/6
        try:
            initial_datas_dict = calculate_detail_excel_document_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
            print_info_calculate_detail_excel_document_sum(detail_excel_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_calculate_detail_excel_document(detail_excel_document_path, e)
            return

        # 计算报销明细 Excel 文档：小表格小计 4/6
        try:
            initial_datas_dict = calculate_detail_excel_document_second_table(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
            print_info_calculate_detail_excel_document_second_table(detail_excel_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_calculate_detail_excel_document(detail_excel_document_path, e)
            return

        # 计算报销明细 Excel 文档：小表格合计 5/6
        try:
            initial_datas_dict = calculate_detail_excel_document_second_table_sum(detail_excel_document_path, initial_datas_dict, sheet_name = "报销明细")
            print_info_calculate_detail_excel_document_second_table_sum(detail_excel_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_calculate_detail_excel_document(detail_excel_document_path, e)
            return

        # 渲染报销明细 Excel 文档 6/6
        try:
            initial_datas_dict = beautify_detail_excel_document(detail_excel_document_path, initial_datas_dict, colors, sheet_name = "报销明细")
            print_info_beautify_detail_excel_document(detail_excel_document_path)
            print_info_detail_excel_document_finished(detail_excel_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_beautify_detail_excel_document(detail_excel_document_path, e)
            return

        # 回写报销明细 Word 文档
        try:
            # 打开文档
            doc = Document(detail_word_document_path)
            # 获取文档中的表格列表（检查是否已被初始化）
            tables = doc.tables
            if tables:

                detail_word_document_add_table(doc, '合计', '', initial_datas_dict['initial_sum'])
                print_info_back_detail_word_document(detail_word_document_path, initial_datas_dict)
            # 保存文档
            doc.save(detail_word_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_back_detail_word_document(detail_word_document_path, e)
            return

        # 搜索出租车情况说明文档
        try:
            if '市内交通' in initial_datas_dict['use_purpose_mapping_datas']:
                taxi_explanation_document_path = find_taxi_explanation_document(target_numeric_folder_path)
                print_info_find_explanation_document('出租车', target_numeric_folder_path, taxi_explanation_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_find_explanation_document('出租车', e)
            return

        # 写入出租车情况说明文档
        try:
            if '市内交通' in initial_datas_dict['use_purpose_mapping_datas']:
                if not destination:
                    print_ten_pentagram()
                    print_error_no_destination()
                    raise ValueError
                write_taxi_explanation_document(taxi_explanation_document_path, detail_word_document_datas, usr_name, destination)
                print_info_write_explanation_document('出租车', taxi_explanation_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_write_explanation_document('出租车', taxi_explanation_document_path, e)
            return

        # 搜索不建固按材料报销情况说明文档
        try:
            # 确保用户输入没有问题（special_material 和 special_material_use 二者成对）
            for special_material, special_material_use in zip(special_materials, special_material_uses):
                special_material_strip = special_material.strip()
                special_material_use_strip = special_material_use.strip()
                if ((special_material_strip is None) and (special_material_use_strip is not None)) or   ((special_material_use_strip is None) and (special_material_strip is not None)):
                    print_ten_pentagram()
                    print_error_special_material_input()
                    raise ValueError

            # special_materials 和 special_material_uses 二者均不为空才操作
            if (not all(special_material is None or str(special_material).strip() == "" for special_material in special_materials)) and (not all(special_material_use is None or str(special_material_use).strip() == "" for special_material_use in special_material_uses)):

                special_material_explanation_document_path = find_special_material_explanation_document(target_numeric_folder_path)
                print_info_find_explanation_document('不建固按材料报销', target_numeric_folder_path, special_material_explanation_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_find_explanation_document('不建固按材料报销', e)
            return

        # 写入不建固按材料报销情况说明文档
        try:
            # 确保用户输入没有问题（special_material 和 special_material_use 二者成对）
            for special_material, special_material_use in zip(special_materials, special_material_uses):
                special_material_strip = special_material.strip()
                special_material_use_strip = special_material_use.strip()
                if ((special_material_strip is None) and (special_material_use_strip is not None)) or ((special_material_use_strip is None) and (special_material_strip is not None)):
                    print_ten_pentagram()
                    print_error_special_material_input()
                    raise ValueError

            # special_materials 和 special_material_uses 二者均不为空才操作
            if (not all(special_material is None or str(special_material).strip() == "" for special_material in special_materials)) and (not all(special_material_use is None or str(special_material_use).strip() == "" for special_material_use in special_material_uses)):

                write_special_material_explanation_document(special_material_explanation_document_path, usr_name, special_materials, special_material_uses)
                print_info_write_explanation_document('不建固按材料报销', special_material_explanation_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_write_explanation_document('不建固按材料报销', special_material_explanation_document_path, e)
            return

    # 每次生成后都要更新后缀
    try:
        updated_detail_word_document_path = update_detail_word_document_suffix(detail_word_document_path)
        print_info_update_detail_word_document_suffix(detail_word_document_path, updated_detail_word_document_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_update_detail_word_document_suffix(detail_word_document_path, e)
        return

    # 输出报销文档
    try:
        zip_folder_path = output_documents(target_numeric_folder_path, output_space)
        print_info_output_documents(output_space, zip_folder_path)
    except Exception as e:
        print_ten_pentagram()
        print_error_output_documents(output_space, e)
        return

# 第三个功能大函数结束

# 第四个功能大函数开始
# 检查报销明细 Word 文档（默认 Microsoft 打开）
def open_detail_word_document(self):
    flag = False
    for item in os.listdir(self.work_space):
        item_path = os.path.join(self.work_space, item)
        # 判断是否是文件夹且文件名是纯数字
        if os.path.isdir(item_path) and item.isdigit():
            flag = True
            break
    if flag:
        # 搜索当前月文件夹
        try:
            target_numeric_folder_path = find_target_numeric_folder(self.work_space)
            print_info_find_target_numeric_folder(self.work_space, target_numeric_folder_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_find_target_numeric_folder(e)
            return

        # 搜索报销明细 Word 文档
        try:
            detail_word_document_path = find_detail_word_document(target_numeric_folder_path)
            print_info_find_detail_word_document(target_numeric_folder_path, detail_word_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_find_detail_word_document(e)
            return

        try:
            os.startfile(detail_word_document_path)  # 在系统默认的应用程序中打开 Word 文档
            print_info_open_detail_word_document(detail_word_document_path)
        except Exception as e:
            print_ten_pentagram()
            print_error_open_detail_word_document(detail_word_document_path, e)
            return
    else:
        main_text_append_error_notfound_files(self)
# 第四个功能大函数结束

# 第五个功能大函数开始
# 在信息表中添加姓名和学号
def save_student_id_in_information_excel_document(information_excel_document_path, student_name, student_id):
    try:
        # 加载 Excel 文件
        workbook = openpyxl.load_workbook(information_excel_document_path)
        ws = workbook.active  # 默认使用第一个工作表

        # 查找第一列为空的行（从第3行开始）
        next_row = None
        for row_index in range(3, ws.max_row + 2):
            if not ws.cell(row=row_index, column=2).value:  # 检查第2列是否为空
                next_row = row_index
                break

        # 如果找到了空行，则写入数据
        if next_row:
            ws.cell(row=next_row, column=2, value=student_name)  # 写入姓名到第2列
            ws.cell(row=next_row, column=3, value=student_id)    # 写入学号到第3列

        # 设置对齐样式：水平和垂直居中
        alignment_style = Alignment(horizontal='center', vertical='center')

        # 设置字体样式
        font_style = Font(name='等线', size=11, bold=False)
   
        # 遍历所有行和列
        for row in ws.iter_rows(min_row=1, min_col=1):
            # 设置行高为16
            ws.row_dimensions[row[0].row].height = 16
            for cell in row:
                if cell.value is not None:  # 如果单元格有内容
                    # 应用对齐样式
                    cell.alignment = alignment_style
                    # 应用字体样式
                    cell.font = font_style

        # 设置列宽
        column_widths = {
            'A': 3,
            'B': 12,
            'C': 20,
        }
        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

        # 保存 Excel 文件
        workbook.save(information_excel_document_path)
        workbook.close()

        print_info_save_student_id_in_information_excel_document(information_excel_document_path, student_name, student_id)
    except Exception as e:
        print_ten_pentagram()
        print_error_save_student_id_in_information_excel_document(information_excel_document_path, e)
        return
# 第五个功能大函数结束

# 第六个功能小函数开始
# 检查目标工作目录是否符合要求（不能直接移动！移一半报错就完蛋了）
def check_new_work_space(initial_work_space, new_work_space):
    ensure_path_exist(new_work_space)
    # 遍历初始工作空间目录下的所有文件夹
    for folder_name in os.listdir(initial_work_space):
        if folder_name:
            folder_path = os.path.join(initial_work_space, folder_name)
            # 判断是否为文件夹且名字为数字
            if os.path.isdir(folder_path) and folder_name.isdigit():
                new_folder_path = os.path.join(new_work_space, folder_name)
                # 如果目标路径已经存在同名文件夹，报错
                if os.path.exists(new_folder_path):
                    raise OSError('目标路径已经存在同名文件夹！')

# 移动到目标工作目录
def to_new_work_space(initial_work_space, new_work_space):
    ensure_path_exist(new_work_space)
    # 遍历初始工作空间目录下的所有文件夹
    for folder_name in os.listdir(initial_work_space):
        folder_path = os.path.join(initial_work_space, folder_name)
        # 判断是否为文件夹且名字为数字
        if os.path.isdir(folder_path) and folder_name.isdigit():
            new_folder_path = os.path.join(new_work_space, folder_name)
            shutil.move(folder_path, new_folder_path)
            print_info_to_new_work_space(folder_path, new_folder_path)
# 第六个功能小函数结束

# 第六个功能大函数开始
# 更换工作目录（移动所有数字文件夹）
def move_to_new_work_space(initial_work_space, new_work_space):
    try:
        check_new_work_space(initial_work_space, new_work_space)
        print_info_check_new_work_space(new_work_space)
    except Exception as e:
        print_ten_pentagram()
        print_error_check_new_work_space(new_work_space, e)
        return
    
    try:
        to_new_work_space(initial_work_space, new_work_space)
        print_info_move_to_new_work_space(new_work_space)
    except Exception as e:
        print_ten_pentagram()
        print_error_move_to_new_work_space(new_work_space, e)
        return
# 第六个功能大函数结束

# 第七个功能小函数开始
# 寻找全部数字文件夹
def find_numeric_folders(work_space):
    # 用于存放符合条件的文件夹路径
    numeric_folder_paths = []

    # 遍历给定路径下的所有文件和文件夹
    for root, dirs, files in os.walk(work_space):
        for dir_name in dirs:
            # 检查文件夹名称是否是由数字组成
            if re.match(r'^\d+$', dir_name):
                # 获取完整路径并添加到结果列表
                numeric_folder_paths.append(os.path.join(root, dir_name))

    return numeric_folder_paths
# 第七个功能小函数结束

# 第七个功能大函数开始
# 搜索
def search(work_space, student_name):
    try:
        numeric_folder_paths = find_numeric_folders(work_space)
        print_info_find_numeric_folders(work_space, numeric_folder_paths)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_numeric_folders(work_space, e)
        return

    # 统计包含信息的报销明细 Word 文档
    try:
        # 每行数据
        all_datas = []
        # 汇总数据
        collected_datas = {}
        for i in range(6):
            collected_datas[use_purpose_number_dict[i]] = 0
        # 开始遍历
        for numeric_folder_path in numeric_folder_paths:
            # 搜索报销明细 Word 文档
            detail_word_document_path = find_detail_word_document(numeric_folder_path)
            # 读取报销明细 Word 文档数据
            detail_word_document_datas = read_detail_word_document(detail_word_document_path)
            # 用于存放符合条件的记录
            matched_detail_word_document_datas = []
            # 遍历所有数组，找到与 student_name 匹配的记录
            for detail_word_document_data in detail_word_document_datas:
                if detail_word_document_data[0] == student_name:
                    matched_detail_word_document_datas.append(detail_word_document_data)
                    for i in range(6):
                        if detail_word_document_data[3].strip() == use_purpose_number_dict[i].strip():
                            collected_datas[use_purpose_number_dict[i]] += float(detail_word_document_data[2])
                            collected_datas[use_purpose_number_dict[i]] = round(collected_datas[use_purpose_number_dict[i]], 2)

            year_month_match = re.search(r"(\d{4}年\d{1,2}月)", os.path.basename(detail_word_document_path))
            if year_month_match:
                year_month = year_month_match.group(0)
                all_datas.append({
                    'year_month': year_month,
                    'data': matched_detail_word_document_datas
                })
        print_info_search(student_name, all_datas, collected_datas)

    except Exception as e:
        print_ten_pentagram()
        print_error_search(e)
        return

# 第七个功能大函数结束

# 第八个功能大函数开始
# 重置
def restart_all(work_space, config_space, clear_all):
    # 重置配置文件夹
    try:
        # 文件路径
        config_path = os.path.join(config_space, "config.json")
        information_excel_document_path = os.path.join(config_space, "information.xlsx")

        if os.path.exists(config_path):
            os.remove(config_path)
        if os.path.exists(information_excel_document_path):
            os.remove(information_excel_document_path)

        startup_config_path(config_path)
        startup_information_excel_document_path(information_excel_document_path)
        print_info_restart_config_space(config_space)

    except Exception as e:
        print_ten_pentagram()
        print_error_restart_config_space(config_space, e)
        return

    # 删除工作目录下文件夹
    try:
        if clear_all:
            numeric_folder_paths = find_numeric_folders(work_space)
            print_info_find_numeric_folders(work_space, numeric_folder_paths)
    except Exception as e:
        print_ten_pentagram()
        print_error_find_numeric_folders(work_space, e)
        return

    try:
        if clear_all:
            for numeric_folder_path in numeric_folder_paths:
                if os.path.exists(numeric_folder_path) and os.path.isdir(numeric_folder_path):
                    # 递归删除文件夹及其内容
                    shutil.rmtree(numeric_folder_path)
                    print_info_del_numeric_folder(numeric_folder_path)
            print_info_restart_work_space(work_space)
    except Exception as e:
        print_ten_pentagram()
        print_error_del_numeric_folder(numeric_folder_path, e)
        return
# 第八个功能大函数结束

# 第九个功能大函数开始
# 出租车目的地检验
def destination_check(work_space):
    try:
        destination_check = False
        destination_student_names = []
        # 搜索当前月文件夹
        target_numeric_folder_path = find_target_numeric_folder(work_space)
        # 搜索报销明细 Word 文档
        detail_word_document_path = find_detail_word_document(target_numeric_folder_path)
        # 报销明细 Word 文档删除合计行
        delete_sum_detail_word_document(detail_word_document_path)
        # 读取报销明细 Word 文档数据
        detail_word_document_datas = read_detail_word_document(detail_word_document_path)
        # 开始遍历
        for detail_word_document_data in detail_word_document_datas:
            if detail_word_document_data[3] == '市内交通':
                destination_check = True
                if detail_word_document_data[0] not in destination_student_names:
                    destination_student_names.append(detail_word_document_data[0])
        return destination_check, destination_student_names
    except Exception as e:
        return False, []
# 第九个功能大函数结束

# PrintStream 类用于重定向标准输出（print）到 GUI 文本框，继承自 QObject
class PrintStream(QObject):
    # 定义信号 signal，用于发送字符串消息
    signal = pyqtSignal(str)
    # write 方法接收消息并通过信号发送给 GUI
    def write(self, message):
        self.signal.emit(message)

# 加载配置文件
def load_config(config_path):
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}

# 写入配置文件
def save_config(config_path, data):
    with open(config_path, 'w', encoding='utf-8') as f:
        # ensure_ascii=False：保存时不将非 ASCII 字符（如中文）转换为 Unicode 转义序列
        # indent=4：指定输出的 JSON 格式化为 4 个空格缩进，方便阅读
        json.dump(data, f, ensure_ascii=False, indent=4)

# 主窗口
# MainWindow 继承自 QMainWindow
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # 主窗口标题
        self.setWindowTitle(f"Easy Accounting {__version__} -- by Mxqwthl")
        self.setGeometry(200, 200, 371, 600)  # 主窗口位置和大小
        # 设置窗口的最小宽度和高度
        self.setMinimumSize(260, 460)

        self.original_stdout = sys.stdout

        # 配置目录
        self.config_space = r"C:\Easy Accounting"
        self.config_path = os.path.join(self.config_space, "config.json")
        self.information_excel_document_path = os.path.join(self.config_space, "information.xlsx")
        # 属性：config 作为配置文件内容
        self.config = load_config(self.config_path)

        # 如果配置文件中不存在这些键，则使用默认值
        self.work_space = self.config.get("work_space", default_config['work_space'])
        self.output_space = self.config.get("output_space", default_config['output_space'])
        self.usr_name = self.config.get("usr_name", default_config['usr_name'])
        self.padding = self.config.get("padding", default_config['padding'])
        self.colors = self.config.get("colors", default_config['colors'])

        # 创建版心
        self.main_widget = QWidget()
        # 将版心设置在中央
        self.setCentralWidget(self.main_widget)

        # 创建主布局【垂直布局】
        self.main_layout = QVBoxLayout()
        self.main_widget.setLayout(self.main_layout) # 添加到版心
        self.main_layout.setSpacing(0)
        self.main_layout.setContentsMargins(10, 10, 10, 10)

        # 第一组：姓名输入【水平布局】
        self.student_name_layout = QHBoxLayout()
        self.main_layout.addLayout(self.student_name_layout) # 添加到主布局【垂直布局】
        self.student_name_layout.setContentsMargins(5, 5, 3, 0)
        # 组件1：标题
        self.student_name_label = QLabel(self)
        self.student_name_label.setObjectName('student_name_label')
        self.student_name_label.setText('姓名：')
        self.student_name_layout.addWidget(self.student_name_label) # 添加到姓名输入【水平布局】
        # 组件2：输入框
        self.student_name_input = QLineEdit(self)
        self.student_name_input.setObjectName('student_name_input')
        self.student_name_input.setPlaceholderText("例：张三") # 提示词
        self.student_name_layout.addWidget(self.student_name_input) # 添加到姓名输入【水平布局】

        self.student_name_label.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Preferred)
        self.student_name_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.student_name_layout.setStretch(0, 0)
        self.student_name_layout.setStretch(1, 1)

        # 第二组：用途输入【水平布局】
        self.use_purpose_layout = QHBoxLayout()
        self.main_layout.addLayout(self.use_purpose_layout) # 添加到主布局【垂直布局】
        self.use_purpose_layout.setContentsMargins(5, 8, 3, 0)
        # 组件1：标题
        self.use_purpose_label = QLabel(self)
        self.use_purpose_label.setObjectName('use_purpose_label')
        self.use_purpose_label.setText('用途：')
        self.use_purpose_layout.addWidget(self.use_purpose_label) # 添加到用途输入【水平布局】
        # 组件2：输入框
        self.use_purpose_input = QLineEdit(self)
        self.use_purpose_input.setObjectName('use_purpose_input')
        self.use_purpose_input.setPlaceholderText("例：3D打印") # 提示词
        self.use_purpose_layout.addWidget(self.use_purpose_input) # 添加到用途输入【水平布局】

        self.use_purpose_label.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Preferred)
        self.use_purpose_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.use_purpose_layout.setStretch(0, 0)
        self.use_purpose_layout.setStretch(1, 1)

        # 第三组：金额输入【水平布局】
        self.money_amount_layout = QHBoxLayout()
        self.main_layout.addLayout(self.money_amount_layout) # 添加到主布局【垂直布局】
        self.money_amount_layout.setContentsMargins(5, 8, 3, 0)
        # 组件1：标题
        self.money_amount_label = QLabel(self)
        self.money_amount_label.setObjectName('money_amount_label')
        self.money_amount_label.setText('金额：')
        self.money_amount_layout.addWidget(self.money_amount_label) # 添加到金额输入【水平布局】
        # 组件2：输入框
        self.money_amount_input = QLineEdit(self)
        self.money_amount_input.setObjectName('money_amount_input')
        self.money_amount_input.setPlaceholderText("例：146.21") # 提示词
        self.money_amount_layout.addWidget(self.money_amount_input) # 添加到金额输入【水平布局】

        self.money_amount_label.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Preferred)
        self.money_amount_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.money_amount_layout.setStretch(0, 0)
        self.money_amount_layout.setStretch(1, 1)

        # 第四组：录入区【水平布局】
        self.clear_and_submit_layout = QHBoxLayout()
        self.main_layout.addLayout(self.clear_and_submit_layout) # 添加到主布局【垂直布局】
        self.clear_and_submit_layout.setContentsMargins(0, 8, 0, 0)
        # 组件1：清除按钮
        self.clear_button = QPushButton('清除', self)
        self.clear_button.setObjectName('clear_button')
        self.clear_button.clicked.connect(self.clear_input)
        self.clear_and_submit_layout.addWidget(self.clear_button)
        # 组件2：录入按钮
        self.submit_button = QPushButton('录入', self)
        self.submit_button.setObjectName('submit_button')
        self.submit_button.clicked.connect(self.run_write_detail_word_document)
        self.clear_and_submit_layout.addWidget(self.submit_button)

        self.clear_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.submit_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.clear_and_submit_layout.setStretch(0, 1)
        self.clear_and_submit_layout.setStretch(1, 1)

        # 第五组：查询区【水平布局】
        self.search_layout = QHBoxLayout()
        self.main_layout.addLayout(self.search_layout) # 添加到主布局【垂直布局】
        self.search_layout.setContentsMargins(3, 3, 3, 0)
        # 组件1：输入框
        self.search_input = QLineEdit(self)
        self.search_input.setObjectName('search_input')
        self.search_input.setPlaceholderText("例：张三") # 提示词
        self.search_layout.addWidget(self.search_input) # 添加到查询区【水平布局】
        # 组件2：查询按钮
        self.search_button = QPushButton('查询', self)
        self.search_button.setObjectName('search_button')
        self.search_button.clicked.connect(self.run_search)
        self.search_layout.addWidget(self.search_button) # 添加到查询区【水平布局】

        self.search_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.search_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.search_layout.setStretch(0, 4)
        self.search_layout.setStretch(1, 1)

        # 第六组：功能区【水平布局】
        self.ribbon_layout = QHBoxLayout()
        self.main_layout.addLayout(self.ribbon_layout) # 添加到主布局【垂直布局】
        self.ribbon_layout.setContentsMargins(0, 8, 0, 0)
        # 组件1：新建按钮
        self.create_button = QPushButton('新建', self)
        self.create_button.setObjectName('create_button')
        self.create_button.clicked.connect(self.run_create_detail_word_document) # 运行 run_create_detail_word_document
        self.ribbon_layout.addWidget(self.create_button) # 添加到功能区【水平布局】
        # 组件2：导出按钮
        self.export_button = QPushButton('导出', self)
        self.export_button.setObjectName('export_button')
        self.export_button.clicked.connect(self.open_export_dialog) # 打开导出子窗口
        self.ribbon_layout.addWidget(self.export_button) # 添加到功能区【水平布局】
        # 组件3：检查按钮
        self.check_button = QPushButton('检查', self)
        self.check_button.setObjectName('check_button')
        self.check_button.clicked.connect(self.run_open_detail_word_document) # 运行 run_open_detail_word_document
        self.ribbon_layout.addWidget(self.check_button) # 添加到功能区【水平布局】

        self.create_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.export_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.check_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.ribbon_layout.setStretch(0, 1)
        self.ribbon_layout.setStretch(1, 1)
        self.ribbon_layout.setStretch(2, 1)

        # 第七组：输出区【单个元素】
        self.main_text = QTextEdit(self)
        self.main_text.setObjectName('main_text')
        self.main_text.setReadOnly(True)
        self.main_layout.addWidget(self.main_text)

        # 第八组：设置区【水平布局】
        self.settings_layout = QHBoxLayout()
        self.main_layout.addLayout(self.settings_layout) # 添加到主布局【垂直布局】
        self.settings_layout.setContentsMargins(0, 8, 0, 0)
        # 组件1：初始设置按钮
        self.initial_settings_button = QPushButton('初始设置', self)
        self.initial_settings_button.setObjectName('initial_settings_button')
        self.initial_settings_button.clicked.connect(self.open_initial_settings_dialog) # 打开初始设置子窗口
        self.settings_layout.addWidget(self.initial_settings_button) # 添加到功能区【水平布局】
        # 组件2：占位符
        self.settings_layout_space = QLabel(self)
        self.settings_layout_space.setObjectName('settings_layout_space')
        self.settings_layout_space.setText('')
        self.settings_layout.addWidget(self.settings_layout_space) # 添加到功能区【水平布局】
        # 组件3：高级设置按钮
        self.advanced_settings_button = QPushButton('高级设置', self)
        self.advanced_settings_button.setObjectName('advanced_settings_button')
        self.advanced_settings_button.clicked.connect(self.open_advanced_settings_dialog) # 打开高级设置子窗口
        self.settings_layout.addWidget(self.advanced_settings_button) # 添加到功能区【水平布局】

        self.initial_settings_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.settings_layout_space.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.advanced_settings_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.settings_layout.setStretch(0, 1)
        self.settings_layout.setStretch(1, 6)
        self.settings_layout.setStretch(2, 1)

        # 重定向 print 输出
        self.print_stream = PrintStream()
        self.print_stream.signal.connect(self.update_output)
        sys.stdout = self.print_stream

        # 初始化
        self.run_startup()

    def run_startup(self):
        startup(self.config_space, self.config_path, self.information_excel_document_path)

    def run_create_detail_word_document(self):
        if self.usr_name.strip() == "":
            self.open_usr_name_input_dialog()
            return
        create_detail_word_document(self.work_space, self.information_excel_document_path, self.usr_name)

    def run_write_detail_word_document(self):
        self.student_name = self.student_name_input.text().strip()
        self.use_purpose = self.use_purpose_input.text().strip()
        self.money_amount = self.money_amount_input.text().strip()

        # 姓名不为空
        if not self.student_name:
            main_text_append_error_input(self, '姓名')
        # 用途不为空
        elif not self.use_purpose:
            main_text_append_error_input(self, '用途')
        elif re.match(r'^\d+(\.\d+)?(、\d+(\.\d+)?)*$', self.money_amount):
            # 金额必须为数字
            if self.usr_name.strip() == "":
                self.open_usr_name_input_dialog()

            self.student_id = find_student_id_by_name(self.information_excel_document_path, self.student_name)
            if self.student_id is None:
                self.open_information_supplementation_dialog()
                return
            
            self.student_id = find_student_id_by_name(self.information_excel_document_path, self.student_name)
            # 防止没输入学号，关闭窗口后，仍然能录入的状况
            if self.student_id:
                if (self.student_id.isdigit() and len(self.student_id) == 10):
                    write_detail_word_document(self.work_space, self.information_excel_document_path, self.usr_name, self.student_name, self.use_purpose, self.money_amount)

                    # 清空录入内容（保留姓名）
                    self.use_purpose_input.setText('')
                    self.money_amount_input.setText('')
        else:
            main_text_append_error_input(self, '金额')

    # 清除输入区
    def clear_input(self):
        self.student_name_input.setText('')
        self.use_purpose_input.setText('')
        self.money_amount_input.setText('')

    # 检查报销明细 Word 文档（默认方式打开）
    def run_open_detail_word_document(self):
        open_detail_word_document(self)

    # 搜索
    def run_search(self):
        self.search_student_name = self.search_input.text().strip()
        # 姓名不为空
        if not self.search_student_name:
            main_text_append_error_input(self, '姓名')
        # restart
        elif self.search_student_name == 'restart':
            self.search_input.setText('')
            self.open_restart_all()
        # 查无此人
        elif not find_student_id_by_name(self.information_excel_document_path, self.search_student_name):
            main_text_append_error_notfound_search_student_name(self)
        else:
            self.search_input.setText('')
            search(self.work_space, self.search_student_name)

    # 打开导出子窗口
    def open_export_dialog(self):
        flag = False
        for item in os.listdir(self.work_space):
            item_path = os.path.join(self.work_space, item)
            # 判断是否是文件夹且文件名是纯数字
            if os.path.isdir(item_path) and item.isdigit():
                flag = True
                break
        if flag:
            # 先获取出租车目的地检验信息
            self.destination_check, self.destination_student_names = destination_check(self.work_space)
            dialog = ExportDialog(self)
            dialog.exec_()
        else:
            main_text_append_error_notfound_files(self)

    # 打开初始设置子窗口
    def open_initial_settings_dialog(self):
        dialog = InitialSettingsDialog(self)
        dialog.exec_()

    # 打开高级设置子窗口
    def open_advanced_settings_dialog(self):
        dialog = AdvancedSettingsDialog(self)
        dialog.exec_()

    # 打开信息补充子窗口
    def open_information_supplementation_dialog(self):
        dialog = InformationSupplementationDialog(self)
        dialog.exec_()

    # 打开用户姓名输入子窗口
    def open_usr_name_input_dialog(self):
        dialog = UsrNameInputDialog(self)
        dialog.exec_()

    # 打开全部初始化子窗口
    def open_restart_all(self):
        dialog = RestartAllDialog(self)
        dialog.exec_()

    # 捕获输出
    def update_output(self, output):
        self.main_text.append(output.rstrip())

    # 恢复标准输出流
    def closeEvent(self, event):
        sys.stdout = self.original_stdout
        super().closeEvent(event)

# QDialog 的特性：
# 1、默认是模态的（阻止用户与其他窗口交互），可以通过 setModal(False) 设置为非模态；

# 2、轻量级，适合短暂显示和用户输入；
# 3、自带一些特有功能，例如 exec_() 方法，可以阻塞主窗口并等待用户操作完成。

# QMainWindow 的特性：
# 1、支持菜单栏、工具栏、状态栏等高级功能；
# 2、更适合作为主窗口，而不是短暂的弹窗。

# 自动定位到父窗口中心
def center_position(parent, self):
    if parent:
        parent_geometry = parent.geometry()
        child_width = self.width()
        child_height = self.height()
        # 计算子窗口的位置，使其显示在父窗口中心
        x = parent_geometry.left() + (parent_geometry.width() - child_width) // 2
        y = parent_geometry.top() + (parent_geometry.height() - child_height) // 2
        self.move(x, y)

class ExportDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # 导出子窗口标题
        self.setWindowTitle("导出")
        # 设置子窗口为模态窗口，阻止用户与其他窗口交互，直到关闭此对话框
        self.setModal(True)
        self.resize(647, 400)  # 主窗口位置和大小
        # 设置窗口的最小宽度和高度
        self.setMinimumSize(647, 400)

        center_position(parent, self)

        # 初始化主布局
        self.main_layout = QVBoxLayout()
        self.setLayout(self.main_layout)  # 直接设置布局到 QDialog
        self.main_layout.setSpacing(0)
        self.main_layout.setContentsMargins(10, 10, 10, 10)

        # Step1
        self.step1 = QTextEdit(self)
        self.step1.setObjectName('step1')
        self.step1.setHtml('<p><strong>Step1：</strong>选择导出目录【所有相关文件将打包发送到此文件夹中】</p>')
        self.step1.setReadOnly(True)  # 设置为只读
        self.step1.setFixedHeight(26)
        self.main_layout.addWidget(self.step1)

        # 第一组：输出路径【水平布局】
        self.output_space_layout = QHBoxLayout()
        self.main_layout.addLayout(self.output_space_layout)
        self.output_space_layout.setContentsMargins(3, 0, 3, 0)
        # 输入 output_space
        self.output_space_input = QLineEdit(self)
        self.output_space_input.setObjectName('output_space_input')
        self.output_space_input.setPlaceholderText('<导出目录>')
        self.output_space_input.setText(self.parent().output_space)
        self.output_space_input.setFixedHeight(36)
        self.output_space_layout.addWidget(self.output_space_input)
        # 输入 output_space 按钮
        self.output_space_button = QPushButton('选择文件夹', self)
        self.output_space_button.setObjectName('work_space_button')
        self.output_space_button.clicked.connect(self.output_space_select)
        self.output_space_button.setFixedHeight(36)
        self.output_space_layout.addWidget(self.output_space_button)

        self.output_space_input.setMinimumWidth(0)  # 允许扩展
        self.output_space_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.output_space_button.setMinimumWidth(0)  # 允许扩展
        self.output_space_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        self.output_space_layout.setStretch(0, 4)
        self.output_space_layout.setStretch(1, 1)

        # Step2
        self.step2 = QTextEdit(self)
        self.step2.setObjectName('step2')
        self.step2.setHtml('<p><strong>Step2：</strong>市内交通相关【据此生成出租车情况说明文档】</p>')
        self.step2.setReadOnly(True)  # 设置为只读
        self.step2.setFixedHeight(36)
        self.main_layout.addWidget(self.step2)

        # 第二组：目的地输入【水平布局】
        self.destination_layout = QVBoxLayout()
        self.main_layout.addLayout(self.destination_layout) # 添加到主布局【垂直布局】
        self.destination_layout.setContentsMargins(3, 0, 3, 0)
        # 组件1：标题
        self.destination_label = QLabel(self)
        if self.parent().destination_student_names:
            destination_label_str = '、'.join(self.parent().destination_student_names) + '的'
        else:
            destination_label_str = ''
        self.destination_label.setText(f'{destination_label_str}出租车目的地：')
        self.destination_label.setObjectName('destination_label')
        self.destination_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.destination_label.setFixedHeight(26)
        self.destination_layout.addWidget(self.destination_label) # 添加到目的地输入【水平布局】
        # 组件2：输入框
        self.destination_input = QTextEdit(self)
        self.destination_input.setObjectName('destination_input')
        self.destination_input.setPlaceholderText("例：启元实验室、友谊宾馆（注意以中文顿号【、】分隔）") # 提示词
        self.destination_input.setFixedHeight(80)
        self.destination_layout.addWidget(self.destination_input) # 添加到目的地输入【水平布局】

        # Step3
        self.step3 = QTextEdit(self)
        self.step3.setObjectName('step3')
        self.step3.setHtml('<p><strong>Step3：</strong>特殊材料相关【据此生成特殊材料情况说明文档】</p>')
        self.step3.setReadOnly(True)  # 设置为只读
        self.step3.setFixedHeight(36)
        self.main_layout.addWidget(self.step3)

        # 第三组：特殊材料 labels【水平布局】
        self.special_material_labels_layout = QHBoxLayout()
        self.main_layout.addLayout(self.special_material_labels_layout) # 添加到主布局【垂直布局】
        self.special_material_labels_layout.setContentsMargins(3, 0, 3, 0)
        # 组件1：特殊材料
        self.special_material_label = QLabel(self)
        self.special_material_label.setText('特殊材料')
        self.special_material_label.setObjectName('special_material_label')
        self.special_material_label.setAlignment(Qt.AlignHCenter)
        self.special_material_label.setFixedHeight(30)
        self.special_material_labels_layout.addWidget(self.special_material_label) # 添加到特殊材料 labels【水平布局】
        # 组件2：用途描述
        self.special_material_use_label = QLabel(self)
        self.special_material_use_label.setText('用途描述')
        self.special_material_use_label.setObjectName('special_material_use_label')
        self.special_material_use_label.setAlignment(Qt.AlignHCenter)
        self.special_material_use_label.setFixedHeight(30)
        self.special_material_labels_layout.addWidget(self.special_material_use_label) # 添加到特殊材料 labels【水平布局】
        # 组件3：添加按钮
        self.add_button = QPushButton('+', self)
        self.add_button.setObjectName('add_button')
        self.add_button.clicked.connect(self.add_material_use_pair)
        self.add_button.setFixedHeight(30)
        self.special_material_labels_layout.addWidget(self.add_button) # 添加到特殊材料 labels【水平布局】
        # 组件4：删除按钮
        self.del_button = QPushButton('-', self)
        self.del_button.setObjectName('del_button')
        self.del_button.clicked.connect(self.del_material_use_pair)
        self.del_button.setFixedHeight(30)
        self.special_material_labels_layout.addWidget(self.del_button) # 添加到特殊材料 labels【水平布局】

        self.special_material_label.setMinimumWidth(0)  # 允许扩展
        self.special_material_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.special_material_use_label.setMinimumWidth(0)  # 允许扩展
        self.special_material_use_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.add_button.setMinimumWidth(0)  # 允许扩展
        self.add_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.del_button.setMinimumWidth(0)  # 允许扩展
        self.del_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        self.special_material_labels_layout.setStretch(0, 6)
        self.special_material_labels_layout.setStretch(1, 6)
        self.special_material_labels_layout.setStretch(2, 1)
        self.special_material_labels_layout.setStretch(3, 1)

        self.special_material_inputs = []
        self.special_material_use_inputs = []
        self.space1s = []
        self.space2s = []

        # 第四组：特殊材料动态输入【水平布局】
        special_material_layout = QHBoxLayout()
        self.main_layout.addLayout(special_material_layout) # 添加到主布局【垂直布局】
        special_material_layout.setContentsMargins(3, 0, 3, 6)
        # 组件1：特殊材料输入
        special_material_input = QTextEdit(self)
        special_material_input.setObjectName('special_material_input')
        special_material_input.setPlaceholderText("例：*电子计算机*迷你电脑（3079.00元）")
        special_material_input.setFixedHeight(55)
        self.special_material_inputs.append(special_material_input)
        special_material_layout.addWidget(special_material_input) # 添加到特殊材料动态输入【水平布局】
        # 组件2：用途描述输入
        special_material_use_input = QTextEdit(self)
        special_material_use_input.setObjectName('special_material_use_input')
        special_material_use_input.setPlaceholderText("例：安装于机器人胸腔，起到控制机器人运动的作用")
        special_material_use_input.setFixedHeight(55)
        self.special_material_use_inputs.append(special_material_use_input)
        special_material_layout.addWidget(special_material_use_input) # 添加到特殊材料动态输入【水平布局】
        # 组件3：占位符1
        space1 = QLabel(self)
        space1.setText('')
        space1.setObjectName('space1')
        self.space1s.append(space1)
        special_material_layout.addWidget(space1) # 添加到特殊材料动态输入【水平布局】
        # 组件4：占位符2
        space2 = QLabel(self)
        space2.setText('')
        space2.setObjectName('space2')
        self.space2s.append(space2)
        special_material_layout.addWidget(space2) # 添加到特殊材料动态输入【水平布局】
    
        special_material_input.setMinimumWidth(0)  # 允许扩展
        special_material_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        special_material_use_input.setMinimumWidth(0)  # 允许扩展
        special_material_use_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        space1.setMinimumWidth(0)  # 允许扩展
        space1.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        space2.setMinimumWidth(0)  # 允许扩展
        space2.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        special_material_layout.setStretch(0, 6)
        special_material_layout.setStretch(1, 6)
        special_material_layout.setStretch(2, 1)
        special_material_layout.setStretch(3, 1)

        # 第五组：确认按钮【单个元素】
        self.confirm_button = QPushButton('确认导出', self)
        self.confirm_button.setObjectName('confirm_button')
        self.confirm_button.clicked.connect(self.confirm) # 运行 confirm
        self.confirm_button.setFixedHeight(48)
        self.main_layout.addWidget(self.confirm_button) # 添加到主布局【垂直布局】

    # 选择导出目录
    def output_space_select(self):
        output_space_selection = QFileDialog.getExistingDirectory(self, '选择目标文件夹')
        output_space_selection = output_space_selection.replace('/', '\\')
        if output_space_selection:
            self.output_space_input.setText(output_space_selection)

    # 添加特殊材料用途描述对
    def add_material_use_pair(self):
        width = self.width()
        height = self.height()
        # 刷新确认按钮
        if self.confirm_button:
            confirm_button = self.confirm_button
            self.main_layout.removeWidget(confirm_button)
            confirm_button.deleteLater()

        # 第四组：特殊材料动态输入【水平布局】
        special_material_layout = QHBoxLayout()
        self.main_layout.addLayout(special_material_layout) # 添加到主布局【垂直布局】
        special_material_layout.setContentsMargins(3, 0, 3, 6)
        # 组件1：特殊材料输入
        special_material_input = QTextEdit(self)
        special_material_input.setObjectName('special_material_input')
        special_material_input.setPlaceholderText('例：*电子计算机*迷你电脑（3079.00元）')
        special_material_input.setFixedHeight(55)
        self.special_material_inputs.append(special_material_input)
        special_material_layout.addWidget(special_material_input) # 添加到特殊材料动态输入【水平布局】
        # 组件2：用途描述输入
        special_material_use_input = QTextEdit(self)
        special_material_use_input.setObjectName('special_material_use_input')
        special_material_use_input.setPlaceholderText('例：安装于机器人胸腔，起到控制机器人运动的作用')
        special_material_use_input.setFixedHeight(55)
        self.special_material_use_inputs.append(special_material_use_input)
        special_material_layout.addWidget(special_material_use_input) # 添加到特殊材料动态输入【水平布局】
        # 组件3：占位符1
        space1 = QLabel(self)
        space1.setText('')
        space1.setObjectName('space1')
        self.space1s.append(space1)
        special_material_layout.addWidget(space1) # 添加到特殊材料动态输入【水平布局】
        # 组件4：占位符2
        space2 = QLabel(self)
        space2.setText('')
        space2.setObjectName('space2')
        self.space2s.append(space2)
        special_material_layout.addWidget(space2) # 添加到特殊材料动态输入【水平布局】

        special_material_input.setMinimumWidth(0)  # 允许扩展
        special_material_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        special_material_use_input.setMinimumWidth(0)  # 允许扩展
        special_material_use_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        space1.setMinimumWidth(0)  # 允许扩展
        space1.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        space2.setMinimumWidth(0)  # 允许扩展
        space2.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        special_material_layout.setStretch(0, 6)
        special_material_layout.setStretch(1, 6)
        special_material_layout.setStretch(2, 1)
        special_material_layout.setStretch(3, 1)

        # 刷新确认按钮
        self.confirm_button = QPushButton('确认导出', self)
        self.confirm_button.setObjectName('confirm_button')
        self.confirm_button.clicked.connect(self.confirm) # 运行 confirm
        self.confirm_button.setFixedHeight(48)
        self.main_layout.addWidget(self.confirm_button) # 添加到主布局【垂直布局】
        self.resize(width, height + 61)
        
    # 删除特殊材料用途描述对
    def del_material_use_pair(self):
        width = self.width()
        height = self.height()

        if self.special_material_inputs and self.special_material_use_inputs:
            # 刷新确认按钮
            if self.confirm_button:
                confirm_button = self.confirm_button
                self.main_layout.removeWidget(confirm_button)
                confirm_button.deleteLater()

            # 从列表中移除最后一组输入框
            special_material_input = self.special_material_inputs.pop()
            special_material_use_input = self.special_material_use_inputs.pop()
            space1 = self.space1s.pop()
            space2 = self.space2s.pop()
            # 从布局中移除对应的组件
            special_material_layout = special_material_input.parentWidget().layout()
            special_material_layout.removeWidget(special_material_input)
            special_material_layout.removeWidget(special_material_use_input)
            special_material_layout.removeWidget(space1)
            special_material_layout.removeWidget(space2)
            # 删除组件以释放资源
            special_material_input.deleteLater()
            special_material_use_input.deleteLater()
            space1.deleteLater()
            space2.deleteLater()
            # 从主布局中移除该水平布局
            self.main_layout.removeItem(special_material_layout)
            del special_material_layout  # 删除水平布局对象以释放资源

            # 刷新确认按钮
            self.confirm_button = QPushButton('确认导出', self)
            self.confirm_button.setObjectName('confirm_button')
            self.confirm_button.clicked.connect(self.confirm) # 运行 confirm
            self.confirm_button.setFixedHeight(48)
            self.main_layout.addWidget(self.confirm_button) # 添加到主布局【垂直布局】

            self.resize(width, height - 61)

    # 确认
    def confirm(self):
        output_space_input = self.output_space_input.text()
        destination_input = self.destination_input.toPlainText()
        special_material_inputs = [input.toPlainText() for input in self.special_material_inputs]
        special_material_use_inputs = [input.toPlainText() for input in self.special_material_use_inputs]
        # 导出工作目录合法
        if not os.path.isabs(output_space_input):
            warning_input(self, '导出目录')
        # 出租车目的地输入根据要求不为空
        elif self.parent().destination_check and not destination_input:
            warning_input(self, '出租车目的地')
        else:
            self.parent().output_space = output_space_input.strip()
            self.parent().destination = destination_input
            self.parent().special_materials = special_material_inputs
            self.parent().special_material_uses = special_material_use_inputs

            update_config(self.parent().config_path, "output_space", self.parent().output_space)

            build_and_output_documents(self.parent().usr_name, self.parent().work_space, self.parent().output_space, self.parent().destination, self.parent().special_materials, self.parent().special_material_uses, self.parent().padding, self.parent().colors)
            self.accept()

class InitialSettingsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # 初始设置子窗口标题
        self.setWindowTitle("初始设置")
        # 设置子窗口为模态窗口，阻止用户与其他窗口交互，直到关闭此对话框
        self.setModal(True)
        self.resize(327, 202)  # 主窗口位置和大小
        # 设置窗口的最小宽度和高度
        self.setMinimumSize(280, 202)

        center_position(parent, self)

        # 初始化主布局
        self.main_layout = QVBoxLayout()
        self.setLayout(self.main_layout)  # 直接设置布局到 QDialog
        self.main_layout.setSpacing(0)
        self.main_layout.setContentsMargins(10, 10, 10, 10)

        # 传递参数
        self.information_excel_document_path = self.parent().information_excel_document_path

        # Step1
        self.step1 = QTextEdit(self)
        self.step1.setObjectName('step1')
        self.step1.setHtml('<p><strong>Step1：</strong>请选择工作目录</p>')
        self.step1.setReadOnly(True)  # 设置为只读
        self.step1.setFixedHeight(26)
        self.main_layout.addWidget(self.step1)

        # 输入 work_space 副布局（横向）
        self.work_space_layout = QHBoxLayout()
        self.main_layout.addLayout(self.work_space_layout)
        self.work_space_layout.setContentsMargins(3, 0, 3, 0)
        # 输入 work_space
        self.work_space_input = QLineEdit(self)
        self.work_space_input.setObjectName('work_space_input')
        self.work_space_input.setPlaceholderText('所有相关文件将创建在此文件夹中')
        self.work_space_input.setFixedHeight(36)
        self.work_space_input.setText(self.parent().work_space)
        # 保存 work_space 初始值
        self.initial_work_space = self.work_space_input.text()
        self.work_space_layout.addWidget(self.work_space_input)
        # 输入 work_space 按钮
        self.work_space_button = QPushButton('选择文件夹', self)
        self.work_space_button.setObjectName('work_space_button')
        self.work_space_button.clicked.connect(self.work_space_select)
        self.work_space_button.setFixedHeight(36)
        self.work_space_layout.addWidget(self.work_space_button)

        self.work_space_input.setMinimumWidth(0)  # 允许扩展
        self.work_space_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.work_space_button.setMinimumWidth(0)  # 允许扩展
        self.work_space_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        self.work_space_layout.setStretch(0, 4)
        self.work_space_layout.setStretch(1, 1)

        # Step2
        self.step2 = QTextEdit(self)
        self.step2.setObjectName('step2')
        self.step2.setHtml('<p><strong>Step2：</strong>请输入用户姓名</p>')
        self.step2.setReadOnly(True)  # 设置为只读
        self.step2.setFixedHeight(36)
        self.main_layout.addWidget(self.step2)

        # 输入 usr_name
        self.usr_name_input = QLineEdit(self)
        self.usr_name_input.setObjectName('usr_name_input')
        self.usr_name_input.setPlaceholderText('例：张三')
        self.usr_name_input.setText(self.parent().usr_name)
        self.usr_name_input.setFixedHeight(36)
        self.main_layout.addWidget(self.usr_name_input)

        # 保存设置按钮
        self.save_initial_settings_button = QPushButton("保存设置")
        self.save_initial_settings_button.setObjectName('save_initial_settings_button')
        self.save_initial_settings_button.clicked.connect(self.save_initial_settings)
        self.save_initial_settings_button.setFixedHeight(48)
        self.main_layout.addWidget(self.save_initial_settings_button)

    def work_space_select(self):
        work_space_selection = QFileDialog.getExistingDirectory(self, '选择目标文件夹')
        work_space_selection = work_space_selection.replace('/', '\\')
        if work_space_selection:
            self.work_space_input.setText(work_space_selection)

    def save_initial_settings(self):
        work_space_input = self.work_space_input.text()
        usr_name_input = self.usr_name_input.text()

        if not os.path.isabs(work_space_input):
            warning_input(self, '工作目录')
        elif not usr_name_input:
            warning_input(self, '用户姓名')
        else:
            self.work_space = work_space_input.strip()
            self.student_name = usr_name_input.strip()
            self.student_id = find_student_id_by_name(self.parent().information_excel_document_path, self.student_name)
            if self.student_id is None:
                self.open_information_supplementation_dialog()
                return
            self.student_id = find_student_id_by_name(self.parent().information_excel_document_path, self.student_name)
            # 防止没输入学号，关闭窗口后，仍然能录入的状况
            if self.student_id:
                if (self.student_id.isdigit() and len(self.student_id) == 10):
                    self.parent().work_space = self.work_space
                    self.parent().usr_name = self.student_name

                    update_config(self.parent().config_path, "work_space", self.parent().work_space)
                    update_config(self.parent().config_path, "usr_name", self.parent().usr_name)

                    # 移动工作目录中的文件夹
                    if os.path.abspath(self.initial_work_space) == os.path.abspath(self.parent().work_space):
                        pass
                    else:
                        move_to_new_work_space(self.initial_work_space, self.parent().work_space)

                    main_text_append_info_setting(self, '工作目录', self.parent().work_space)
                    main_text_append_info_setting(self, '用户姓名', self.parent().usr_name)
                    self.accept()

    # 打开信息补充子窗口
    def open_information_supplementation_dialog(self):
        dialog = InformationSupplementationDialog(self)
        dialog.exec_()

class AdvancedSettingsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # 高级设置子窗口标题
        self.setWindowTitle("高级设置")
        # 设置子窗口为模态窗口，阻止用户与其他窗口交互，直到关闭此对话框
        self.setModal(True)
        self.setFixedSize(270, 300)

        center_position(parent, self)

        # 从主窗口传入参数（因为有恢复默认而不保存设置，防止改变主窗口状态）【关键】
        self.padding = self.parent().padding
        self.colors = self.parent().colors

        # 初始化主布局【垂直布局】
        self.main_layout = QVBoxLayout()
        self.setLayout(self.main_layout)  # 直接设置布局到 QDialog
        self.main_layout.setSpacing(0)
        self.main_layout.setContentsMargins(10, 10, 10, 10)

        # 第一组：padding 设置【单个元素】
        self.padding_checkbox = QCheckBox('Padding', self)
        self.padding_checkbox.setObjectName('padding_checkbox')
        self.padding_checkbox.setChecked(self.padding)
        self.main_layout.addWidget(self.padding_checkbox)

        # 第二组：颜色设置【水平布局】
        self.color_inputs = []
        self.number_to_use = {
            0: '材料',
            1: '市内交通',
            2: '物流',
            3: '打印',
            4: '差旅',
            5: '论文',
        }
        
        for i in range(6):
            # 第二组：颜色设置【水平布局】
            color_layout = QHBoxLayout()
            self.main_layout.addLayout(color_layout) # 添加到主布局【垂直布局】
            color_layout.setContentsMargins(0, 10, 0, 0)
            # 组件1：标题
            color_label = QLabel(self)
            color_label.setText(self.number_to_use.get(i, 'error'))
            color_label.setObjectName('color_label')
            color_label.setAlignment(Qt.AlignCenter)
            color_layout.addWidget(color_label) # 添加到颜色设置【水平布局】
            # 组件2：颜色输入
            color_input = QLineEdit(self)
            color_input.setObjectName('color_input')
            color_input.setText(f'{self.colors[i]}')
            color_input.setStyleSheet(f'background-color: {self.colors[i]};')
            self.color_inputs.append(color_input)
            color_layout.addWidget(color_input) # 添加到颜色设置【水平布局】
            # 绑定输入框的 textChanged 信号
            color_input.textChanged.connect(lambda text, input=color_input: self.update_background_color(text, input))
            # 组件3：颜色选择
            color_button = QPushButton('选择颜色', self)
            color_button.setObjectName('color_button')
            color_button.clicked.connect(lambda _, input=color_input: self.select_color(input))
            color_layout.addWidget(color_button) # 添加到颜色设置【水平布局】

            color_label.setMinimumWidth(0)  # 允许扩展
            color_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
            color_input.setMinimumWidth(0)  # 允许扩展
            color_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
            color_button.setMinimumWidth(0)  # 允许扩展
            color_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

            color_layout.setStretch(0, 1)
            color_layout.setStretch(1, 1)
            color_layout.setStretch(2, 1)

        # 第三组：保存设置【水平布局】
        self.save_settings_layout = QHBoxLayout()
        self.main_layout.addLayout(self.save_settings_layout) # 添加到主布局【垂直布局】
        self.save_settings_layout.setContentsMargins(0, 10, 0, 0)
        # 组件1：恢复默认按钮
        self.restore_default_settings_button = QPushButton('恢复默认', self)
        self.restore_default_settings_button.setObjectName('restore_default_settings_button')
        self.restore_default_settings_button.clicked.connect(self.restore_default_settings)
        self.save_settings_layout.addWidget(self.restore_default_settings_button) # 添加到保存设置【水平布局】
        # 组件2：占位符
        self.save_settings_layout_space = QLabel(self)
        self.save_settings_layout_space.setText('')
        self.save_settings_layout_space.setObjectName('save_settings_layout_space')
        self.save_settings_layout_space.setAlignment(Qt.AlignCenter)
        self.save_settings_layout.addWidget(self.save_settings_layout_space) # 添加到保存设置【水平布局】
        # 组件3：保存设置按钮
        self.save_advanced_settings_button = QPushButton('保存设置', self)
        self.save_advanced_settings_button.setObjectName('save_advanced_settings_button')
        self.save_advanced_settings_button.clicked.connect(self.save_advanced_settings)
        self.save_settings_layout.addWidget(self.save_advanced_settings_button) # 添加到保存设置【水平布局】

        self.restore_default_settings_button.setMinimumWidth(0)  # 允许扩展
        self.restore_default_settings_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.save_settings_layout_space.setMinimumWidth(0)  # 允许扩展
        self.save_settings_layout_space.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.save_advanced_settings_button.setMinimumWidth(0)  # 允许扩展
        self.save_advanced_settings_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        self.save_settings_layout.setStretch(0, 1)
        self.save_settings_layout.setStretch(1, 1)
        self.save_settings_layout.setStretch(2, 1)

    def select_color(self, input_field):
        color = QColorDialog.getColor()
        if color.isValid():
            input_field.setText(color.name().upper())

    def update_background_color(self, text, input):
        # 验证输入是否是有效的十六进制颜色代码
        if self.is_valid_hex_color(text):
            # 动态更新输入框背景颜色
            input.setStyleSheet(f"background-color: {text};")
        else:
            # 如果输入无效，恢复默认背景颜色
            input.setStyleSheet("background-color: #ffffff;")

    def is_valid_hex_color(self, text):
        if text.startswith("#") and len(text) == 7:
            try:
                int(text[1:], 16)  # 尝试将十六进制代码转换为整数
                return True
            except ValueError:
                pass
        return False

    # 保存设置
    def save_advanced_settings(self):
        
        self.padding = self.padding_checkbox.isChecked()
        self.colors = [color_input.text().strip() for color_input in self.color_inputs]
        for color in self.colors:
            if not self.is_valid_hex_color(color):
                warning_input(self, '颜色')
        
        color_check = True
        for color in self.colors:
            if not self.is_valid_hex_color(color):
                color_check = False
        if color_check:
            self.parent().padding = self.padding
            self.parent().colors = self.colors

            update_config(self.parent().config_path, "padding", self.parent().padding)
            update_config(self.parent().config_path, "colors", self.parent().colors)

            main_text_append_info_setting(self, 'padding', self.parent().padding)
            for i in range(6):
                main_text_append_info_setting(self, '颜色设置 - ' + self.number_to_use.get(i, 'error'), self.parent().colors[i])
            self.accept()
    
    # 恢复默认
    def restore_default_settings(self):
        # 子窗口属性变化
        self.padding = default_config['padding']
        self.colors = default_config['colors']
        # 子窗口显示变化
        self.padding_checkbox.setChecked(self.padding)
        for i, color_input in enumerate(self.color_inputs):
            if i < len(self.colors):  # 确保不会超出索引范围
                color_input.setText(f'{self.colors[i]}')  # 设置新的颜色值
        # 配置文件不能变化！！！（如果用户不点保存）

class InformationSupplementationDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # 信息补充子窗口标题
        self.setWindowTitle("信息补充")
        # 设置子窗口为模态窗口，阻止用户与其他窗口交互，直到关闭此对话框
        self.setModal(True)
        self.setFixedSize(250, 125)

        center_position(parent, self)

        # 初始化主布局【垂直布局】
        self.main_layout = QVBoxLayout()
        self.setLayout(self.main_layout)  # 直接设置布局到 QDialog
        self.main_layout.setSpacing(0)
        self.main_layout.setContentsMargins(10, 10, 10, 10)

        # 第一组：学号输入标签【单个元素】
        self.student_id_label = QLabel(self)
        self.student_id_label.setObjectName('student_id_label')
        self.student_id_label.setText(f'请输入【{self.parent().student_name}】的学号：')
        self.main_layout.addWidget(self.student_id_label) # 添加到主布局【垂直布局】
        # 第二组：学号输入【单个元素】
        self.student_id_input = QLineEdit(self)
        self.student_id_input.setObjectName('student_id_input')
        self.student_id_input.setPlaceholderText("例：1120191884")
        self.main_layout.addWidget(self.student_id_input) # 添加到主布局【垂直布局】
        # 第三组：保存按钮【单个元素】
        self.student_id_save_button = QPushButton('保存', self)
        self.student_id_save_button.setObjectName('student_id_save_button')
        self.student_id_save_button.clicked.connect(self.run_save_student_id_in_information_excel_document)
        self.main_layout.addWidget(self.student_id_save_button) # 添加到主布局【垂直布局】

    def run_save_student_id_in_information_excel_document(self):
        self.parent().student_id = self.student_id_input.text().strip()

        # 验证学号
        if not (self.parent().student_id.isdigit() and len(self.parent().student_id) == 10):
            # 显示错误提示
            warning_input(self, '学号')
        if (self.parent().student_id.isdigit() and len(self.parent().student_id) == 10):
            save_student_id_in_information_excel_document(self.parent().information_excel_document_path, self.parent().student_name, self.parent().student_id)
            self.accept()

class UsrNameInputDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # 用户姓名输入子窗口标题
        self.setWindowTitle("用户姓名输入")
        # 设置子窗口为模态窗口，阻止用户与其他窗口交互，直到关闭此对话框
        self.setModal(True)
        self.setFixedSize(292, 125)

        center_position(parent, self)
        # 传递参数
        self.information_excel_document_path = self.parent().information_excel_document_path

        # 初始化主布局【垂直布局】
        self.main_layout = QVBoxLayout()
        self.setLayout(self.main_layout)  # 直接设置布局到 QDialog
        self.main_layout.setSpacing(0)
        self.main_layout.setContentsMargins(10, 10, 10, 10)

        # 第一组：用户姓名输入标签【单个元素】
        self.usr_name_label = QLabel(self)
        self.usr_name_label.setObjectName('usr_name_label')
        self.usr_name_label.setText(f'未检测到用户姓名！请输入用户姓名：')
        self.main_layout.addWidget(self.usr_name_label) # 添加到主布局【垂直布局】
        # 第二组：用户姓名输入【单个元素】
        self.usr_name_input = QLineEdit(self)
        self.usr_name_input.setObjectName('usr_name_input')
        self.usr_name_input.setPlaceholderText("例：张三")
        self.usr_name_input.setStyleSheet('margin-top: 10px')
        self.main_layout.addWidget(self.usr_name_input) # 添加到主布局【垂直布局】
        # 第三组：保存按钮【单个元素】
        self.usr_name_save_button = QPushButton('保存', self)
        self.usr_name_save_button.setObjectName('usr_name_save_button')
        self.usr_name_save_button.clicked.connect(self.save_usr_name)
        self.usr_name_save_button.setStyleSheet('margin-top: 10px')
        self.main_layout.addWidget(self.usr_name_save_button) # 添加到主布局【垂直布局】
    
    def save_usr_name(self):
        usr_name_input = self.usr_name_input.text()
        if usr_name_input is None:
            warning_input(self, '用户姓名')
        if usr_name_input:
            self.student_name = usr_name_input.strip()
            self.student_id = find_student_id_by_name(self.parent().information_excel_document_path, self.student_name)
            if self.student_id is None:
                self.open_information_supplementation_dialog()
                return
            if (self.student_id.isdigit() and len(self.student_id) == 10):
                self.parent().usr_name = self.student_name
                update_config(self.parent().config_path, "usr_name", self.parent().usr_name)
                main_text_append_info_setting(self, '用户姓名', self.parent().usr_name)
                self.accept()

    # 打开信息补充子窗口
    def open_information_supplementation_dialog(self):
        dialog = InformationSupplementationDialog(self)
        dialog.exec_()

# 重置子窗口
class RestartAllDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # 重置子窗口标题
        self.setWindowTitle("重置")
        # 设置子窗口为模态窗口，阻止用户与其他窗口交互，直到关闭此对话框
        self.setModal(True)
        self.setFixedSize(139, 87)

        center_position(parent, self)
        
        # 初始化主布局【垂直布局】
        self.main_layout = QVBoxLayout()
        self.setLayout(self.main_layout)  # 直接设置布局到 QDialog
        self.main_layout.setSpacing(0)
        self.main_layout.setContentsMargins(10, 10, 10, 10)

        # 第一组：clear_all【单个元素】
        self.clear_all_checkbox = QCheckBox('清除工作目录', self)
        self.clear_all_checkbox.setObjectName('clear_all_checkbox')
        self.clear_all_checkbox.setChecked(False)
        self.main_layout.addWidget(self.clear_all_checkbox)
        # 第二组：重置按钮【单个元素】
        self.restart_all_button = QPushButton('重置', self)
        self.restart_all_button.setObjectName('restart_all_button')
        self.restart_all_button.clicked.connect(self.run_restart_all)
        self.main_layout.addWidget(self.restart_all_button) # 添加到主布局【垂直布局】

    def run_restart_all(self):
        self.clear_all = self.clear_all_checkbox.isChecked()
        restart_all(self.parent().work_space, self.parent().config_space, self.clear_all)
        self.accept()

# 程序主入口
def main():
    app = QApplication(sys.argv)
    app.setStyleSheet(load_stylesheet('style.qss'))
    app.setWindowIcon(QIcon(resource_path("favicon.ico")))
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()